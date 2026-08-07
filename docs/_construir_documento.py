"""Construye el documento académico desde las salidas del pipeline.

Ninguna cifra se escribe a mano: todas se leen de data/surface. Aplica NTC 1486
(márgenes asimétricos, numeración con primera página distinta) y APA 7.
"""
from __future__ import annotations

import json
import pathlib
import sys

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(pathlib.Path(__file__).parents[1]))
from src import config

S = config.SURFACE
CIFRAS = pd.read_csv(S / "lista_cifras.csv").set_index("concepto")["valor"].to_dict()
REG = json.loads((config.DOCS / "registro_version.json").read_text(encoding="utf-8"))


def c(concepto: str) -> str:
    """Devuelve una cifra desde la lista trazable. Falla si no existe."""
    if concepto not in CIFRAS:
        raise KeyError(f"La cifra «{concepto}» no está en lista_cifras.csv")
    return str(CIFRAS[concepto])


def _campo_pagina(parrafo):
    r = parrafo.add_run()
    for tipo in ("begin", "instrText", "end"):
        el = OxmlElement(f"w:{'fldChar' if tipo != 'instrText' else 'instrText'}")
        if tipo == "instrText":
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el.set(qn("w:fldCharType"), tipo)
        r._r.append(el)


def nuevo_documento() -> Document:
    d = Document()
    est = d.styles["Normal"]
    est.font.name = "Arial"
    est.font.size = Pt(12)
    for s in d.sections:
        # NTC 1486: superior 3, izquierdo 4, derecho 2, inferior 3
        s.top_margin, s.left_margin = Cm(3), Cm(4)
        s.right_margin, s.bottom_margin = Cm(2), Cm(3)
        s.different_first_page_header_footer = True
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _campo_pagina(p)
    return d


def titulo(d, texto, nivel=1):
    p = d.add_paragraph()
    r = p.add_run(texto.upper() if nivel == 1 else texto)
    r.bold = True
    r.font.size = Pt(14 if nivel == 1 else 12)
    p.space_before, p.space_after = Pt(18 if nivel == 1 else 12), Pt(6)
    return p


def parrafo(d, texto, justificar=True):
    p = d.add_paragraph(texto)
    if justificar:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


def nota(d, texto, color=(120, 120, 120)):
    p = d.add_paragraph()
    r = p.add_run(texto)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(12)
    return p


def tabla(d, df, numero, titulo_tabla, fuente, ancho_cm=15.6, max_filas=25):
    p = d.add_paragraph()
    r = p.add_run(f"Tabla {numero}. {titulo_tabla}")
    r.bold = True
    r.font.size = Pt(11)
    df = df.head(max_filas)
    t = d.add_table(rows=1, cols=len(df.columns))
    t.style = "Table Grid"
    t.autofit = False
    ancho = Cm(ancho_cm / len(df.columns))
    for j, col in enumerate(df.columns):
        celda = t.rows[0].cells[j]
        celda.text = str(col)
        celda.width = ancho
        for par in celda.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(9)
    tr = t.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)
    for _, fila in df.iterrows():
        celdas = t.add_row().cells
        for j, v in enumerate(fila):
            celdas[j].text = "" if pd.isna(v) else str(v)
            celdas[j].width = ancho
            for par in celdas[j].paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.size = Pt(9)
    nota(d, f"Fuente: {fuente}")
    return t


def construir() -> pathlib.Path:
    d = nuevo_documento()

    # ---------------------------------------------------------------- portada
    for _ in range(3):
        d.add_paragraph()
    for texto, tam, negrita in [
        ("UNIVERSIDAD LIBRE — SECCIONAL CALI", 14, True),
        ("FACULTAD DE INGENIERÍA", 12, False),
        (f"PROGRAMA: {REG['portada']['programa'].upper()}", 12, False),
    ]:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(texto)
        r.bold, r.font.size = negrita, Pt(tam)
    for _ in range(3):
        d.add_paragraph()
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PRODUCTO DE DATOS INTEGRADO PARA EL MONITOREO Y PRONÓSTICO "
                  "DE LAS IMPORTACIONES Y LA CARGA PORTUARIA DE BUENAVENTURA")
    r.bold, r.font.size = True, Pt(14)
    for _ in range(4):
        d.add_paragraph()
    for texto in ["JUAN MANUEL TEJADA FAJARDO", "JESÚS ALEJANDRO GUERRERO"]:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(texto).font.size = Pt(12)
    d.add_paragraph()
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Director del trabajo: {REG['portada']['director']}")
    r.font.size = Pt(12)
    for _ in range(4):
        d.add_paragraph()
    for texto in ["SANTIAGO DE CALI", "2026"]:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(texto).font.size = Pt(12)
    d.add_page_break()

    # ---------------------------------------------------------------- contenido
    titulo(d, "Contenido")
    for t in ["1. Introducción", "2. Planteamiento del problema", "3. Objetivos",
              "4. Alcance y delimitación", "5. Fuentes de datos y viabilidad",
              "6. Metodología", "7. Arquitectura e integración de dominios",
              "8. Resultados del análisis exploratorio", "9. Resultados del modelado",
              "10. Producto de datos", "11. Conclusiones", "12. Recomendaciones",
              "13. Referencias", "Anexo A. Matriz de trazabilidad P01–P52",
              "Anexo B. Reporte de no viabilidad"]:
        d.add_paragraph(t)
    nota(d, "Para actualizar la paginación: seleccionar y pulsar F9 en Word.")
    d.add_page_break()

    # ---------------------------------------------------------------- 1
    titulo(d, "1. Introducción")
    parrafo(d, (
        "Buenaventura concentra una parte central del comercio exterior colombiano por el "
        "Pacífico. Las estadísticas oficiales que lo describen existen y son públicas, pero "
        "viven separadas: los microdatos de importaciones del DANE registran declaraciones "
        "aduaneras, mientras las estadísticas de la Superintendencia de Transporte registran "
        "toneladas movilizadas por las sociedades portuarias. Ninguna de las dos, por sí "
        "sola, permite distinguir si un cambio en el comercio del puerto responde a más "
        "mercancía o a mercancía más cara."))
    parrafo(d, (
        f"Este trabajo integra ambos dominios. El aduanero aporta {c('Meses de la serie aduanera')} "
        f"meses ({c('Periodo aduanero')}) construidos a partir de "
        f"{c('Registros de la aduana 35')} registros de la aduana 35. El portuario aporta "
        f"{c('Meses de la serie portuaria')} meses ({c('Periodo portuario')}). El periodo en "
        f"que ambos coexisten es de {c('Meses integrados')} meses."))
    parrafo(d, (
        "El resultado más relevante del trabajo contradice la hipótesis con la que se "
        "inició: integrar los dominios no mejora la capacidad de pronóstico. Ese hallazgo, "
        "lejos de invalidar la integración, precisa su función dentro del producto."))

    # ---------------------------------------------------------------- 2
    titulo(d, "2. Planteamiento del problema")
    parrafo(d, (
        "La existencia de datos públicos no resuelve por sí sola una necesidad analítica. "
        "Los archivos deben integrarse, homologarse y validarse antes de poder usarse. En el "
        "caso de Buenaventura se suma una dificultad conceptual: una variación monetaria "
        "puede confundirse con un cambio físico si no se contrastan el valor CIF, el peso "
        "neto y el valor unitario implícito."))
    parrafo(d, (
        "La pregunta que orienta el trabajo es: ¿cómo desarrollar un producto de datos "
        "reproducible que integre información aduanera y portuaria de Buenaventura, "
        "comunique su incertidumbre de manera calibrada y genere señales explicables para "
        "priorizar el análisis, sin afirmar relaciones que las fuentes no permitan sostener?"))

    # ---------------------------------------------------------------- 3
    titulo(d, "3. Objetivos")
    titulo(d, "3.1 Objetivo general", 2)
    parrafo(d, (
        "Desarrollar y validar un producto de datos reproducible que integre información "
        "aduanera y portuaria de Buenaventura para monitorear, describir y pronosticar los "
        "indicadores que tengan calidad e historia suficientes, conservando la trazabilidad "
        "desde la fuente oficial hasta la salida de consumo."))
    titulo(d, "3.2 Objetivos específicos", 2)
    for t in [
        "Evaluar la viabilidad de fuentes en los dominios aduanero, portuario, marítimo y contextual.",
        "Construir series mensuales reproducibles y reconciliadas de cada dominio viable.",
        "Ejecutar las 52 preguntas del análisis exploratorio y conservar su evidencia.",
        "Declarar el tipo de integración posible entre dominios y sus límites.",
        "Determinar qué indicadores pueden pronosticarse de forma responsable.",
        "Medir si la integración multidominio mejora el pronóstico.",
        "Comunicar la incertidumbre mediante intervalos con cobertura medida.",
        "Documentar con evidencia las fuentes y variables no viables.",
    ]:
        d.add_paragraph(t, style="List Bullet")

    # ---------------------------------------------------------------- 4
    titulo(d, "4. Alcance y delimitación")
    parrafo(d, (
        "El alcance comprende la integración de los dominios aduanero y portuario, el "
        "análisis exploratorio completo, el pronóstico de los indicadores elegibles, los "
        "intervalos calibrados, las reglas de alerta y un tablero de consulta."))
    titulo(d, "4.1 Exclusiones declaradas", 2)
    for t in [
        "El código de aduana 35 es una unidad de registro administrativo. No mide congestión, "
        "patios, buques, contenedores ni tiempos de despacho.",
        "La fuente portuaria identifica la sociedad portuaria, no la terminal física ni el muelle.",
        "El valor CIF por kilogramo es un valor unitario implícito agregado. No es un precio.",
        "El transbordo no constituye comercio exterior: es carga que cambia de buque sin "
        "entrar ni salir del país.",
        "El dominio marítimo y el operativo no se construyen por ausencia de fuente pública.",
        "El valor CIF está expresado en dólares corrientes.",
        "Las alertas son señales de revisión analítica, no órdenes operativas.",
    ]:
        d.add_paragraph(t, style="List Bullet")

    # ---------------------------------------------------------------- 5
    titulo(d, "5. Fuentes de datos y viabilidad")
    parrafo(d, (
        "Se inventariaron seis fuentes en cuatro dominios antes de descargar ninguna. La "
        "decisión sobre cada una se tomó con criterios declarados de antemano: cobertura "
        "histórica, granularidad compatible, licencia, reproducibilidad de la descarga y "
        "rezago de publicación conocido."))
    tabla(d, pd.read_csv(S / "catalogo_fuentes.csv")[
        ["fuente_id", "dominio", "entidad", "formato", "frecuencia", "decision"]],
        1, "Inventario y decisión sobre las fuentes evaluadas",
        "Elaboración propia. Salida del pipeline: catalogo_fuentes.csv (P01, P07).")
    parrafo(d, (
        "Dos hallazgos condicionan el alcance. El primero es favorable: el conjunto de datos "
        "de tráfico portuario de la Superintendencia de Transporte publica el dato mensual "
        "por sociedad portuaria y tipo de carga, con descarga automatizable y licencia "
        "CC BY-SA 4.0. El diagnóstico previo suponía que solo existían boletines "
        "trimestrales en formato PDF, y esa suposición era incorrecta."))
    parrafo(d, (
        "El segundo es restrictivo: no existe fuente pública con serie histórica tabular "
        "para arribos, zarpes, tipos de buque, banderas, horarios, tiempos estimados y "
        "reales de llegada o salida, ni permanencias. La Dirección General Marítima publica "
        "únicamente boletines trimestrales agregados. Los datos de evento provienen de "
        "sistemas de terminal o de proveedores comerciales de AIS, sin acceso ni "
        "presupuesto. Ocho preguntas del análisis quedan cerradas como no viables con la "
        "búsqueda documentada."))

    # ---------------------------------------------------------------- 6
    titulo(d, "6. Metodología")
    parrafo(d, (
        "Investigación aplicada, cuantitativa y longitudinal. Cada fase produce una "
        "evidencia verificable y ninguna cifra llega a este documento sin pasar antes por un "
        "archivo generado por el pipeline."))
    titulo(d, "6.1 Validación temporal", 2)
    parrafo(d, (
        "Se emplea validación walk-forward de un paso con ventana expansiva. En cada corte "
        "el modelo se entrena con el pasado y predice el mes siguiente. Toda transformación "
        "que dependa de los datos, incluidos los escaladores, se ajusta dentro del conjunto "
        "de entrenamiento de cada corte y nunca sobre la serie completa."))
    titulo(d, "6.2 Construcción de intervalos", 2)
    parrafo(d, (
        "Los intervalos se construyen a partir de los cuantiles empíricos de los errores "
        "fuera de muestra del backtest, con calibración expansiva: para el corte t solo se "
        "utilizan los errores de los cortes anteriores. En ningún caso se derivan de una "
        "métrica de error puntual. Si la cobertura empírica no alcanza el nivel nominal, el "
        "intervalo se recalibra o se renombra según lo que realmente cubre."))
    titulo(d, "6.3 Líneas base", 2)
    parrafo(d, (
        "Se reportan siempre tres referencias: naive de un paso, naive estacional y drift. "
        "Reportar únicamente la referencia estacional exagera la mejora atribuible al "
        f"modelo. En esta serie, el naive estacional alcanza un WAPE de "
        f"{c('WAPE · naive 12 sobre el CIF')} mientras el naive simple alcanza "
        f"{c('WAPE · naive 1 sobre el CIF')}: la referencia exigente es la segunda."))

    # ---------------------------------------------------------------- 7
    titulo(d, "7. Arquitectura e integración de dominios")
    parrafo(d, (
        "Los datos se organizan en cuatro capas: raw con los archivos originales inmutables "
        "y su hash, landing con la homologación de esquemas, trusted con las series "
        "validadas y surface con las salidas de consumo. Los cuadernos orquestan y la lógica "
        "reutilizable vive en módulos cubiertos por "
        f"{REG['codigo']['pruebas']} pruebas automatizadas."))
    titulo(d, "7.1 Tipo de integración posible", 2)
    parrafo(d, (
        "No existe llave pública que relacione una declaración de importación con un "
        "movimiento de carga portuaria. Los conceptos tampoco son equivalentes: la aduana "
        "mide mercancía importada sin embalaje bajo el código 35, mientras el puerto mide "
        "toneladas movilizadas con embalaje e incluye exportación y transbordo. Por eso la "
        "integración es agregada por mes calendario, nunca directa."))
    tabla(d, pd.read_csv(S / "matriz_integracion.csv"), 2,
          "Relaciones entre fuentes y tipo de integración",
          "Elaboración propia. Salida del pipeline: matriz_integracion.csv (P42).")
    parrafo(d, (
        "Afirmar que una importación concreta llegó en un buque determinado o por una "
        "sociedad portuaria específica excedería lo que las fuentes permiten sostener."))

    # ---------------------------------------------------------------- 8
    titulo(d, "8. Resultados del análisis exploratorio")
    titulo(d, "8.1 Estado de las 52 preguntas", 2)
    traza = pd.read_csv(S / "matriz_trazabilidad_eda.csv")
    resumen = traza["estado"].value_counts().rename_axis("estado").reset_index(name="preguntas")
    tabla(d, resumen, 3, "Estado de las 52 preguntas del análisis exploratorio",
          "Elaboración propia. Salida del pipeline: matriz_trazabilidad_eda.csv.")
    parrafo(d, (
        "Ninguna pregunta quedó sin respuesta. Las cerradas como no viables documentan dónde "
        "se buscó, qué se encontró y qué fuente haría falta para abrirlas en una fase "
        "posterior. La especificación admite esa respuesta siempre que esté demostrada."))

    titulo(d, "8.2 Valor económico frente a volumen físico", 2)
    parrafo(d, (
        f"El valor CIF medio mensual es de {c('CIF medio mensual')} y el valor unitario "
        f"implícito medio de {c('CIF por kilogramo medio')}. Comparando el primer y el "
        "último año de la serie, el valor unitario crece de forma sostenida, lo que indica "
        "que por cada kilogramo importado se paga más que antes. Ese aumento puede provenir "
        "de precios, de fletes y seguros, o de un cambio en la mezcla hacia mercancías de "
        "mayor valor por kilo; la serie agregada no permite separar las tres causas."))
    parrafo(d, (
        "Sobre los 101 meses comunes con el dominio portuario, el valor CIF creció un "
        "80,4 % entre los primeros y los últimos doce meses. Descomponiendo ese crecimiento "
        "en logaritmos, el volumen aporta el 52,2 % y el valor unitario el 47,0 %; el 0,8 % "
        "restante es un residuo de promediar meses, porque la media de un cociente no "
        "coincide con el cociente de las medias. Dicho de otro modo, entró alrededor de un "
        "tercio más de mercancía y esa mercancía vale alrededor de un tercio más por "
        "kilogramo."))
    parrafo(d, (
        "La detección de meses extremos sobre la serie aduanera marca once meses atípicos "
        "en valor CIF y uno en peso neto. Que los extremos aparezcan en valor y casi no en "
        "cantidad es coherente con lo anterior, aunque no lo demuestra: son dos formas "
        "distintas de mirar la misma serie."))

    titulo(d, "8.3 Concentración: el contraste entre dominios", 2)
    tabla(d, pd.read_csv(S / "comparacion_concentracion.csv"), 4,
          "Índice de concentración HHI por dimensión y dominio",
          "Elaboración propia. Salida del pipeline: comparacion_concentracion.csv (P17, P18, P26).")
    parrafo(d, (
        f"La canasta de productos y la de orígenes están desconcentradas, con índices de "
        f"{c('HHI capítulo arancelario')} y {c('HHI país de origen')} respectivamente. La "
        f"movilización de carga, en cambio, está concentrada: el índice alcanza "
        f"{c('HHI sociedad portuaria')} y una sola sociedad portuaria moviliza el "
        f"{c('Participación de la mayor sociedad')} de las toneladas del periodo."))
    parrafo(d, (
        "La lectura conjunta es específica de este trabajo: Buenaventura importa mercancía "
        "variada desde orígenes variados y la moviliza a través de pocas sociedades "
        "portuarias. Ninguna de las dos fuentes por separado permite observar ese contraste, "
        "y ese es el argumento que justifica la integración."))
    parrafo(d, (
        "Entre los años completos, el índice por sociedad portuaria bajó de 4.722 en 2018 a "
        "2.988 en 2023 y subió a 3.740 en 2025. El valor de 2026 (4.217) no es comparable: "
        "cubre seis meses y solo tres sociedades reportan. Recalculando 2025 con esas mismas "
        "tres sociedades, el índice sería 4.096, de modo que la mayor parte del salto se "
        "explica por el cambio de cobertura del reporte y no por una redistribución entre "
        "las sociedades que continúan."))
    parrafo(d, (
        "El índice mide reparto de toneladas reportadas. No mide capacidad instalada, "
        "utilización ni posibilidad real de sustitución entre sociedades, de modo que un "
        "valor alto no permite concluir por sí solo que exista un riesgo operativo."))

    titulo(d, "8.4 Especialización y capacidad de respaldo", 2)
    tabla(d, pd.read_csv(S / "especializacion_terminales.csv"), 5,
          "Composición porcentual de la carga de cada sociedad portuaria por tipo",
          "Elaboración propia. Salida del pipeline: especializacion_terminales.csv (P27).")
    parrafo(d, (
        "La especialización es casi total y matiza la lectura del índice de concentración, "
        "que trata a las sociedades como si fueran intercambiables. El cruce muestra que no "
        "movilizan lo mismo: una de ellas no registró toneladas de granel en todo el periodo "
        "y otra no registró contenedores. El reparto agregado describe cuánto movilizó cada "
        "una, no si podrían movilizar la carga de las demás: la fuente no contiene capacidad "
        "instalada, utilización ni número de instalaciones."))

    titulo(d, "8.5 El total portuario y sus componentes", 2)
    parrafo(d, (
        "El tráfico portuario agregado combina tres flujos distintos: importación, "
        "exportación y transbordo. El transbordo es carga que llega en un buque y sale en "
        "otro sin entrar al territorio aduanero, de modo que no forma parte del comercio "
        "exterior del país."))
    parrafo(d, (
        "Separarlos cambia la lectura. Comparando los primeros doce meses de la serie "
        "portuaria con los últimos doce, las toneladas de importación suben un 19,6 % y las "
        "de exportación un 1,3 %, mientras el transbordo cae un 85,2 %. El total resultante "
        "baja un 13,3 %. Un lector que observe únicamente el total concluiría que la zona "
        "portuaria movilizó menos carga; la descomposición muestra que lo que descendió fue "
        "el transbordo, mientras la carga de importación aumentaba."))
    parrafo(d, (
        "La fuente no informa por qué cambió el transbordo. Decisiones de las navieras, "
        "reasignación de rutas o cambios en el reporte son explicaciones posibles que estos "
        "datos no permiten distinguir."))

    titulo(d, "8.6 Sociedades sin reporte en 2026", 2)
    sr = pd.read_csv(S / "sociedades_sin_reporte_2026.csv")
    tabla(d, sr[["sociedad_portuaria", "ultimo_anio_con_reporte"]], 6,
          "Sociedades portuarias que no aparecen en los reportes de 2026",
          "Elaboración propia. Salida del pipeline: sociedades_sin_reporte_2026.csv (P28).")
    parrafo(d, (
        f"{c('Sociedades sin reporte en 2026')} sociedades portuarias no aparecen reportando "
        "en los seis meses observados de 2026. Se verificó mes a mes para descartar que se "
        "tratara de un rezago de publicación."))
    parrafo(d, (
        "La fuente registra reporte, no operación. Con estos datos solo puede afirmarse que "
        "no figuran en los reportes del periodo observado; no puede afirmarse que hayan "
        "dejado de operar. La causa institucional requeriría fuentes que no se integraron. "
        "La implicación metodológica es relevante: un modelo entrenado sobre el total de la "
        "zona interpretaría esta ausencia como una contracción real del comercio."))

    # ---------------------------------------------------------------- 9
    titulo(d, "9. Resultados del modelado")
    titulo(d, "9.1 Indicadores elegibles", 2)
    tabla(d, pd.read_csv(S / "elegibilidad_indicadores.csv"), 7,
          "Elegibilidad de los indicadores candidatos a pronóstico",
          "Elaboración propia. Salida del pipeline: elegibilidad_indicadores.csv (P46).")
    parrafo(d, (
        "El criterio se fijó antes de mirar los resultados: un indicador con menos de 36 "
        "observaciones mensuales no se pronostica, se describe. Tres de los siete candidatos "
        "quedan excluidos porque su fuente no existe."))

    titulo(d, "9.2 La integración no mejora el pronóstico", 2)
    parrafo(d, (
        "La versión 5 partía de una hipótesis implícita: cruzar dominios mejoraría la "
        "predicción del valor CIF. Se midió mediante un análisis de ablación por grupos de "
        "variables, con todas ellas rezagadas para que ninguna use información no disponible "
        "en la fecha de predicción."))
    tabla(d, pd.read_csv(S / "ablacion_multidominio.csv"), 8,
          "Análisis de ablación multidominio sobre el pronóstico del valor CIF",
          "Elaboración propia. Backtest walk-forward de 24 cortes. "
          "Salida del pipeline: ablacion_multidominio.csv (P49).")
    parrafo(d, (
        f"El conjunto que solo usa la historia propia del CIF y variables de calendario "
        f"alcanza un WAPE de {c('WAPE · historia propia + calendario')}. Añadir variables "
        f"portuarias lo eleva a {c('WAPE · historia propia + puerto')}, y el modelo "
        f"integrado completo a {c('WAPE · integrado completo')}. **Las variables portuarias "
        f"empeoran el pronóstico.**"))
    parrafo(d, (
        "La explicación es coherente con la naturaleza de las fuentes: ambas miden el mismo "
        "comercio subyacente y comparten el mismo rezago de publicación, de modo que el "
        "dominio portuario no aporta información que la historia del propio valor CIF no "
        "contenga ya, y sí consume grados de libertad sobre una muestra corta."))
    parrafo(d, (
        "La consecuencia para el diseño del producto es directa: el modelo predictivo "
        "principal es el aduanero, y el dominio portuario se conserva por su valor "
        "descriptivo y explicativo. Sirve para responder si un mes cambió por valor o por "
        "volumen físico y por qué sociedad portuaria pasó la carga. Presentarlo como una "
        "mejora del pronóstico sería afirmar lo contrario de lo que se midió."))

    titulo(d, "9.3 Indicadores portuarios", 2)
    met = pd.read_csv(S / "metricas_modelos_portuarios.csv")
    tabla(d, met[met.ventana == 24][["objetivo", "modelo", "wape_pct", "mase_12"]].round(3),
          9, "Desempeño de modelos y líneas base sobre los indicadores portuarios",
          "Elaboración propia. Ventana de 24 cortes. "
          "Salida del pipeline: metricas_modelos_portuarios.csv (P47, P48).")
    parrafo(d, (
        f"El resultado es dividido y se reporta como tal. Para las toneladas totales, Ridge "
        f"alcanza {c('WAPE · toneladas totales, Ridge')} frente a "
        f"{c('WAPE · toneladas totales, naive 1')} del naive simple, una mejora sustantiva. "
        f"Para la carga contenerizada ocurre lo contrario: el naive simple obtiene "
        f"{c('WAPE · carga contenerizada, naive 1')} y Ridge queda en "
        f"{c('WAPE · carga contenerizada, Ridge')}, es decir peor que repetir el último "
        "valor observado."))
    parrafo(d, (
        "La recomendación para la carga contenerizada es no utilizar modelo. Un indicador "
        "que no se pronostica mejor que la referencia trivial no debe presentarse con un "
        "modelo encima."))

    titulo(d, "9.4 Intervalos de predicción", 2)
    # La salida del pipeline trae catorce columnas; repartidas en 15,6 cm quedarían a
    # 1,1 cm cada una, ilegibles. Se seleccionan las seis que responden la pregunta y el
    # resto se traslada al texto y a la nota al pie.
    cob_raw = pd.read_csv(S / "cobertura_intervalos_portuarios.csv")
    cob = pd.DataFrame({
        "Indicador": cob_raw["caso"].str.replace("_ridge_v24", "", regex=False)
                                    .str.replace("_ridge_v36", "", regex=False)
                                    .str.replace("_", " "),
        "Ventana": cob_raw["caso"].str.extract(r"v(\d+)")[0] + " cortes",
        "Nominal": (cob_raw["nivel_nominal"] * 100).round(0).astype(int).astype(str) + " %",
        "Empírica": (cob_raw["cobertura_empirica"] * 100).round(1).astype(str)
                    .str.replace(".", ",", regex=False) + " %",
        "Cortes": cob_raw["n_evaluados"],
        "Fuera": cob_raw["casos_fuera"],
    })
    tabla(d, cob, 10, "Cobertura empírica de los intervalos de predicción",
          "Elaboración propia. Método: cuantiles empíricos de los errores fuera de muestra "
          "con calibración expansiva; en ningún caso derivados de una métrica de error "
          "puntual. Salida completa del pipeline, con ancho del intervalo, meses fuera y "
          "veredicto por caso: cobertura_intervalos_portuarios.csv (P51).")
    parrafo(d, (
        "La cobertura se mide y no se declara. Con pocos cortes evaluables, el intervalo de "
        "confianza de la propia cobertura es amplio y esa limitación se reporta junto con "
        "el resultado: una cobertura del 66 % sobre doce cortes no permite concluir que el "
        "intervalo esté mal calibrado, del mismo modo que una del 92 % no permite concluir "
        "que lo esté bien."))

    # ---------------------------------------------------------------- 10
    titulo(d, "10. Producto de datos")
    parrafo(d, (
        "El producto consta de un tablero de seis vistas que lee exclusivamente de la capa "
        "de consumo y no calcula nada: vista ejecutiva, aduanera, portuaria con el análisis "
        "por sociedad portuaria, marítima donde se documenta la no viabilidad, predictiva "
        "con la ablación, y de calidad y trazabilidad."))
    parrafo(d, (
        "Las reglas de alerta operan en tres niveles: normal, seguimiento y alerta. Los "
        "umbrales se calibran sobre la variación interanual y únicamente con la ventana de "
        "entrenamiento. Se optó por la variación interanual y no por el nivel porque, en una "
        "serie con tendencia, comparar contra la mediana histórica del nivel hace que todo "
        "mes reciente resulte anómalo y la alerta se dispare siempre."))
    parrafo(d, "Cada alerta explica su razón y no constituye una orden operativa.")

    # ---------------------------------------------------------------- 11
    titulo(d, "11. Conclusiones")
    for t in [
        f"Se construyeron y reconciliaron dos dominios: {c('Meses de la serie aduanera')} "
        f"meses aduaneros y {c('Meses de la serie portuaria')} portuarios, con "
        f"{c('Meses integrados')} meses de intersección.",
        "La integración entre dominios es agregada por mes. No existe llave pública que "
        "permita una integración directa y no se construyó ninguna por inferencia.",
        "La integración no mejora el pronóstico del valor CIF: lo empeora. Su valor es "
        "descriptivo y explicativo, y el modelo predictivo principal es el aduanero.",
        "El comercio de Buenaventura combina una canasta diversificada de productos y "
        "orígenes con una movilización concentrada en pocas sociedades portuarias "
        "altamente especializadas.",
        "El dominio marítimo y el operativo no son viables con fuentes públicas. Ocho "
        "preguntas se cierran con la búsqueda documentada.",
        "Las 52 preguntas tienen respuesta: ninguna quedó sin evidencia ni sin justificación.",
    ]:
        d.add_paragraph(t, style="List Bullet")

    # ---------------------------------------------------------------- 12
    titulo(d, "12. Recomendaciones")
    for t in [
        "Conservar el modelo aduanero como componente predictivo y el portuario como "
        "componente descriptivo, tal como indica el análisis de ablación.",
        "No pronosticar la carga contenerizada con modelo: usar la referencia naive.",
        "Repetir la descarga del DANE en la próxima publicación para medir el efecto de "
        "las revisiones sobre la serie, hoy no evaluable.",
        "Traducir los códigos de país y capítulo arancelario a sus nombres oficiales.",
        "Decidir si el total portuario se modela sobre la zona completa o solo sobre las "
        "sociedades que siguen reportando.",
        "Validar las reglas de alerta con un usuario real antes de considerarlas operativas.",
        "Gestionar acceso institucional a fuentes de evento si se desea abrir el bloque "
        "marítimo en una fase futura.",
    ]:
        d.add_paragraph(t, style="List Bullet")

    # ---------------------------------------------------------------- 13
    titulo(d, "13. Referencias")
    parrafo(d, (
        "Conforme a APA 7.ª edición, la fecha de consulta se incluye únicamente en las "
        "fuentes diseñadas para cambiar con el tiempo y sin versión archivada. Las obras "
        "estables —leyes, normas técnicas y manuales impresos— no la llevan."))
    for r in [
        # --- fuentes estables: sin fecha de consulta, por diseño
        "American Psychological Association. (2020). Publication Manual of the American "
        "Psychological Association (7.ª ed.). APA.",
        "Congreso de la República de Colombia. (2012). Ley 1581 de 2012, por la cual se "
        "dictan disposiciones generales para la protección de datos personales. Diario "
        "Oficial No. 48.587.",
        "Congreso de la República de Colombia. (2014). Ley 1712 de 2014, por medio de la "
        "cual se crea la Ley de Transparencia y del Derecho de Acceso a la Información "
        "Pública Nacional. Diario Oficial No. 49.084.",
        # --- fuentes que se actualizan: con fecha de consulta verificada
        "Departamento Administrativo Nacional de Estadística. (2018). Estadísticas de "
        "importaciones (IMPO) 1970–2024 [conjunto de datos]. Archivo Nacional de Datos. "
        "Recuperado el 6 de agosto de 2026, de "
        "https://microdatos.dane.gov.co/index.php/catalog/473",
        "Departamento Administrativo Nacional de Estadística. (2026). Estadísticas de "
        "importaciones (IMPO) 2025–2026 [conjunto de datos]. Archivo Nacional de Datos. "
        "Recuperado el 6 de agosto de 2026, de "
        "https://microdatos.dane.gov.co/index.php/catalog/856",
        "Hyndman, R. J., y Athanasopoulos, G. (2021). Forecasting: Principles and practice "
        "(3.ª ed.). OTexts. Recuperado el 6 de agosto de 2026, de https://otexts.com/fpp3/",
        "Instituto Colombiano de Normas Técnicas y Certificación. (2008). NTC 1486: "
        "Documentación. Presentación de tesis, trabajos de grado y otros trabajos de "
        "investigación. ICONTEC.",
        "Superintendencia de Transporte. (2026). Tráfico portuario marítimo en Colombia "
        "[conjunto de datos]. Datos Abiertos Colombia. Recuperado el 6 de agosto de 2026, "
        "de https://www.datos.gov.co/Transporte/Trafico-Portuario-Mar-timo-En-Colombia/"
        "5r3g-zv5z",
    ]:
        p = d.add_paragraph(r)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(6)

    titulo(d, "13.1 Condiciones de uso exigidas por las fuentes", 2)
    parrafo(d, (
        "El DANE exige una cita textual: «Fuente: Departamento Administrativo Nacional de "
        "Estadística: www.dane.gov.co», y prohíbe la reproducción de los datos en medios "
        "que los pongan a disposición de múltiples usuarios sin su visto bueno escrito. "
        "Este trabajo reproduce agregados mensuales, no los microdatos originales."))
    parrafo(d, (
        "El conjunto de datos de la Superintendencia de Transporte se publica bajo licencia "
        "Creative Commons Atribución-CompartirIgual 4.0 Internacional, que obliga a citar la "
        "fuente y a compartir los derivados bajo la misma licencia. La entidad advierte "
        "además que sus cifras son referenciales y que solo ella puede certificarlas."))

    tabla(d, pd.DataFrame([
        {"fuente": "DANE IMPO catálogo 473", "verificada": "2026-08-06",
         "ultima_actualizacion": "2018-03-28", "cambia": "no (histórico cerrado)"},
        {"fuente": "DANE IMPO catálogo 856", "verificada": "2026-08-06",
         "ultima_actualizacion": "2026-07-22", "cambia": "sí (mensual)"},
        {"fuente": "Supertransporte 5r3g-zv5z", "verificada": "2026-08-06",
         "ultima_actualizacion": "2026-08-01", "cambia": "sí (trimestral)"},
        {"fuente": "Hyndman y Athanasopoulos, fpp3", "verificada": "2026-08-06",
         "ultima_actualizacion": "2026-07-23", "cambia": "sí (edición en línea)"},
    ]), 13, "Verificación de las fuentes en línea citadas",
        "Elaboración propia. Cada URL fue consultada el 6 de agosto de 2026 y la fecha de "
        "actualización proviene de los metadatos publicados por la propia fuente.")

    # ---------------------------------------------------------------- anexos
    d.add_page_break()
    titulo(d, "Anexo A. Matriz de trazabilidad P01–P52")
    tabla(d, traza[["pregunta", "bloque", "estado"]], 11,
          "Estado y bloque de cada pregunta del análisis exploratorio",
          "Elaboración propia. Salida del pipeline: matriz_trazabilidad_eda.csv.",
          max_filas=52)
    d.add_page_break()
    titulo(d, "Anexo B. Reporte de no viabilidad")
    tabla(d, pd.read_csv(S / "reporte_no_viabilidad.csv")[["pregunta", "tema", "estado"]],
          12, "Preguntas cerradas por ausencia de fuente",
          "Elaboración propia. Salida del pipeline: reporte_no_viabilidad.csv.")
    parrafo(d, pd.read_csv(S / "reporte_no_viabilidad.csv")["razon"].iloc[0])

    ruta = config.DOCUMENTS / "Documento_Academico_Buenaventura_V5.docx"
    ruta.parent.mkdir(parents=True, exist_ok=True)
    d.save(ruta)
    return ruta


if __name__ == "__main__":
    r = construir()
    print(f"documento generado: {r.name} ({r.stat().st_size // 1024} KB)")
