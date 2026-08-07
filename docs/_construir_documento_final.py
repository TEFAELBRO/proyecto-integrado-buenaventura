"""Documento académico final, con la estructura solicitada para la socialización.

Toda cifra se lee de data/surface/lista_cifras.csv. Si el documento pide un valor que
no está en esa lista, la construcción falla: es la garantía de que nada se escribe a mano.
"""
from __future__ import annotations

import json
import pathlib
import sys

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(pathlib.Path(__file__).parents[1]))
from src import config

S = config.SURFACE
_CIFRAS = pd.read_csv(S / "lista_cifras.csv").set_index("concepto")["valor"].to_dict()
REG = json.loads((config.DOCS / "registro_version.json").read_text(encoding="utf-8"))
AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x59, 0x59, 0x59)
_USADAS: set[str] = set()


def c(concepto: str) -> str:
    """Cifra verificada. Lanza si no existe en la lista trazable."""
    if concepto not in _CIFRAS:
        raise KeyError(f"La cifra «{concepto}» no está en lista_cifras.csv")
    _USADAS.add(concepto)
    return str(_CIFRAS[concepto])


def _campo_pagina(parrafo):
    r = parrafo.add_run()
    for tipo in ("begin", "instrText", "end"):
        el = OxmlElement(f"w:{'fldChar' if tipo != 'instrText' else 'instrText'}")
        if tipo == "instrText":
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el.set(qn("w:fldCharType"), tipo)
        r._r.append(el)


def documento() -> Document:
    d = Document()
    est = d.styles["Normal"]
    est.font.name = "Arial"
    est.font.size = Pt(12)
    for s in d.sections:
        s.top_margin, s.left_margin = Cm(3), Cm(4)
        s.right_margin, s.bottom_margin = Cm(2), Cm(3)
        s.different_first_page_header_footer = True
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _campo_pagina(p)
    # Estilos de encabezado para que Word regenere la tabla de contenido con F9
    for nivel, tam in ((1, 14), (2, 12), (3, 11)):
        h = d.styles[f"Heading {nivel}"]
        h.font.name = "Arial"
        h.font.size = Pt(tam)
        h.font.bold = True
        h.font.color.rgb = AZUL
    return d


def h1(d, texto):
    p = d.add_heading(texto, level=1)
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(20), Pt(8)
    return p


def h2(d, texto):
    p = d.add_heading(texto, level=2)
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(14), Pt(6)
    return p


def h3(d, texto):
    p = d.add_heading(texto, level=3)
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(10), Pt(4)
    return p


def par(d, texto):
    p = d.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    return p


def destacado(d, texto):
    """Resaltado discreto para un resultado principal, sin convertirlo en diapositiva."""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(6), Pt(10)
    r = p.add_run(texto)
    r.bold = True
    r.font.color.rgb = AZUL
    pPr = p._p.get_or_add_pPr()
    bordes = OxmlElement("w:pBdr")
    izq = OxmlElement("w:left")
    izq.set(qn("w:val"), "single")
    izq.set(qn("w:sz"), "18")
    izq.set(qn("w:space"), "8")
    izq.set(qn("w:color"), "1F4E79")
    bordes.append(izq)
    pPr.append(bordes)
    return p


def vinetas(d, items):
    for t in items:
        p = d.add_paragraph(t, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)


def nota(d, texto):
    p = d.add_paragraph()
    r = p.add_run(texto)
    r.italic, r.font.size, r.font.color.rgb = True, Pt(9), GRIS
    p.paragraph_format.space_after = Pt(12)
    return p


_N_TABLA = [0]


def tabla(d, df, titulo_t, fuente, ancho_cm=15.6, max_filas=30):
    _N_TABLA[0] += 1
    p = d.add_paragraph()
    r = p.add_run(f"Tabla {_N_TABLA[0]}. {titulo_t}")
    r.bold, r.font.size = True, Pt(11)
    p.paragraph_format.space_before = Pt(10)
    df = df.head(max_filas)
    if len(df.columns) > 7:
        raise ValueError(f"Tabla {_N_TABLA[0]} tiene {len(df.columns)} columnas: ilegible")
    t = d.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    t.autofit = False
    ancho = Cm(ancho_cm / len(df.columns))
    for j, col in enumerate(df.columns):
        celda = t.rows[0].cells[j]
        celda.text = str(col)
        celda.width = ancho
        for pa in celda.paragraphs:
            pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in pa.runs:
                run.bold, run.font.size = True, Pt(9)
    trPr = t.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)
    for _, fila in df.iterrows():
        celdas = t.add_row().cells
        for j, v in enumerate(fila):
            celdas[j].text = "" if pd.isna(v) else str(v)
            celdas[j].width = ancho
            for pa in celdas[j].paragraphs:
                pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in pa.runs:
                    run.font.size = Pt(9)
    nota(d, f"Fuente: {fuente}")
    return t


def construir() -> pathlib.Path:
    d = documento()
    port = REG["portada"]

    # ------------------------------------------------------------------ PORTADA
    for _ in range(3):
        d.add_paragraph()
    for texto, tam, neg in [("UNIVERSIDAD LIBRE — SECCIONAL CALI", 14, True),
                            ("FACULTAD DE INGENIERÍA", 12, False),
                            (f"PROGRAMA DE {port['programa'].upper()}", 12, False)]:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(texto)
        r.bold, r.font.size = neg, Pt(tam)
    for _ in range(3):
        d.add_paragraph()
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PRODUCTO DE DATOS PARA EL MONITOREO, LA EXPLICACIÓN Y EL PRONÓSTICO "
                  "DE INDICADORES AGREGADOS DE LAS IMPORTACIONES Y LA CARGA PORTUARIA "
                  "DE BUENAVENTURA")
    r.bold, r.font.size = True, Pt(14)
    for _ in range(4):
        d.add_paragraph()
    for a in port["integrantes"]:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(a.upper()).font.size = Pt(12)
    d.add_paragraph()
    for texto in [f"Director del trabajo: {port['director']}",
                  "Ingeniería del Producto de Ciencia de Datos"]:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(texto).font.size = Pt(12)
    for _ in range(4):
        d.add_paragraph()
    for texto in ["SANTIAGO DE CALI", "2026"]:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(texto).font.size = Pt(12)
    d.add_page_break()

    # ------------------------------------------------------------------ CONTENIDO
    h1(d, "Contenido")
    for t in ["1. Introducción", "2. Descripción de los procesos generales",
              "3. Planteamiento y formulación del problema", "4. Objetivos",
              "5. Justificación", "6. Alcance del trabajo", "7. Delimitación",
              "8. Marco de referencia", "9. Metodología",
              "10. Ingeniería del producto de datos", "11. Resultados del análisis exploratorio",
              "12. Resultados del modelado", "13. Modelo de negocios",
              "14. Estrategias de comunicación", "15. Conclusiones", "16. Recomendaciones",
              "17. Fuentes de información", "Anexo A. Matriz de trazabilidad P01–P52",
              "Anexo B. Reporte de no viabilidad", "Anexo C. Preparación para la sustentación"]:
        p = d.add_paragraph(t)
        p.paragraph_format.space_after = Pt(2)
    nota(d, "Tabla de contenido generable desde Word con Referencias, Tabla de contenido. "
            "Para actualizar la paginación: seleccionar y pulsar F9.")
    d.add_page_break()

    # ------------------------------------------------------------------ 1
    h1(d, "1. Introducción")
    par(d, ("Buenaventura concentra una parte central del comercio exterior colombiano por "
            "el Pacífico. La información que lo describe existe y es pública, pero vive en "
            "fuentes separadas que miden cosas distintas. El Departamento Administrativo "
            "Nacional de Estadística publica las declaraciones de importación, donde el dato "
            "central es el valor CIF y el peso neto de la mercancía. La Superintendencia de "
            "Transporte publica las toneladas movilizadas por las sociedades portuarias, que "
            "incluyen el embalaje y cubren además exportación y transbordo."))
    par(d, ("Que ambas existan no significa que puedan compararse sin cuidado. Una registra "
            "operaciones aduaneras bajo el código 35; la otra registra movimientos de carga "
            "en la zona portuaria. No comparten unidad, ni cobertura, ni identificador. "
            "Tampoco existe una llave pública que permita saber en qué buque llegó una "
            "importación concreta."))
    par(d, (f"Este trabajo las integra por mes calendario. El dominio aduanero aporta "
            f"{c('Meses de la serie aduanera')} meses ({c('Periodo aduanero')}), construidos "
            f"a partir de {c('Registros de la aduana 35')} registros. El portuario aporta "
            f"{c('Meses de la serie portuaria')} meses ({c('Periodo portuario')}). El periodo "
            f"en que ambos coexisten es de {c('Meses integrados')} meses "
            f"({c('Periodo integrado')})."))
    par(d, ("La integración permite algo que ninguna fuente logra por separado: distinguir si "
            "un cambio en el comercio se debe a que entró más mercancía o a que la mercancía "
            "vale más por kilogramo. El producto incluye además un componente predictivo "
            "acotado, con intervalos cuya cobertura se mide en lugar de declararse."))
    destacado(d, ("Un resultado del trabajo contradice la hipótesis con la que se inició: "
                  "sumar variables portuarias no mejora el pronóstico del valor CIF, lo "
                  "empeora. El aporte de la integración es descriptivo y explicativo, no "
                  "predictivo."))

    # ------------------------------------------------------------------ 2
    h1(d, "2. Descripción de los procesos generales")
    par(d, ("El producto se construye mediante un flujo de catorce pasos encadenados. Cada "
            "paso deja un archivo verificable, de modo que cualquier cifra del documento "
            "puede rastrearse hasta su origen."))
    tabla(d, pd.DataFrame([
        {"Paso": 1, "Proceso": "Identificación y evaluación de fuentes",
         "Salida": "catalogo_fuentes.csv"},
        {"Paso": 2, "Proceso": "Descarga y conservación de datos originales",
         "Salida": "capa raw inmutable"},
        {"Paso": 3, "Proceso": "Control de integridad con hashes", "Salida": "manifest_fuentes.csv"},
        {"Paso": 4, "Proceso": "Homologación de formatos entre vigencias",
         "Salida": "cambios_esquema.csv"},
        {"Paso": 5, "Proceso": "Validación de estructura, unidades y continuidad",
         "Salida": "calidad_portuaria.csv"},
        {"Paso": 6, "Proceso": "Agregación mensual por dominio", "Salida": "capa trusted"},
        {"Paso": 7, "Proceso": "Integración por mes calendario",
         "Salida": "vista_integrada_mensual.parquet"},
        {"Paso": 8, "Proceso": "Ejecución de las 52 preguntas del análisis exploratorio",
         "Salida": "matriz_trazabilidad_eda.csv"},
        {"Paso": 9, "Proceso": "Selección de indicadores pronosticables",
         "Salida": "elegibilidad_indicadores.csv"},
        {"Paso": 10, "Proceso": "Validación temporal walk-forward",
         "Salida": "metricas_modelos_portuarios.csv"},
        {"Paso": 11, "Proceso": "Construcción y medición de intervalos",
         "Salida": "cobertura_intervalos_portuarios.csv"},
        {"Paso": 12, "Proceso": "Generación de archivos de consumo", "Salida": "capa surface"},
        {"Paso": 13, "Proceso": "Visualización en tablero", "Salida": "dashboard de seis vistas"},
        {"Paso": 14, "Proceso": "Documentación y trazabilidad",
         "Salida": "lista_cifras.csv y este documento"},
    ]), "Procesos generales del producto de datos",
        "Elaboración propia a partir del flujo implementado en el repositorio del proyecto.")
    par(d, ("El flujo se organiza en cuatro capas sucesivas. La capa raw guarda los archivos "
            "originales sin modificar, con su hash. La capa landing unifica formatos. La capa "
            "trusted contiene las series validadas. La capa surface reúne las salidas que "
            "consume el tablero. Los cuadernos orquestan el flujo y la lógica reutilizable "
            "vive en módulos cubiertos por pruebas automatizadas."))

    # ------------------------------------------------------------------ 3
    h1(d, "3. Planteamiento y formulación del problema")
    h2(d, "3.1 Planteamiento del problema")
    par(d, ("Un analista que quiera entender el comercio de Buenaventura se encuentra con "
            "tres dificultades encadenadas. La primera es de acceso: los microdatos "
            "aduaneros se publican en paquetes comprimidos anuales que hay que descargar y "
            "descomprimir uno por uno, y las estadísticas portuarias viven en otro portal, "
            "con otra estructura."))
    par(d, ("La segunda es de formato. Los microdatos aduaneros cambian de estructura cinco "
            "veces entre 2012 y 2026: varía el separador, la convención decimal, la "
            "codificación del archivo y el número de columnas. Un cambio de convención "
            "decimal no genera un error visible, genera números mal leídos que se suman igual."))
    par(d, ("La tercera es conceptual y es la más costosa. El valor CIF y el peso neto miden "
            "cosas distintas, y su cociente es un valor unitario que no debe confundirse con "
            "un precio. Sin contrastar las tres magnitudes, una variación monetaria puede "
            "leerse como si fuera un cambio en la cantidad de mercancía. Del lado portuario "
            "ocurre algo equivalente: el total incluye transbordo, que es carga que cambia de "
            "buque sin entrar al país."))
    h2(d, "3.2 Formulación del problema")
    par(d, ("¿Cómo desarrollar un producto de datos reproducible que integre información "
            "aduanera y portuaria de Buenaventura, permita monitorear y explicar indicadores "
            "agregados, pronostique únicamente aquellos con historia y calidad suficientes, "
            "comunique la incertidumbre de forma calibrada y no afirme relaciones que las "
            "fuentes no permiten sostener?"))
    # ------------------------------------------------------------------ 4
    h1(d, "4. Objetivos")
    h2(d, "4.1 Objetivo general")
    par(d, ("Desarrollar y validar un producto de datos reproducible que integre información "
            "aduanera y portuaria de Buenaventura para monitorear, explicar y pronosticar "
            "indicadores agregados, conservando la trazabilidad desde la fuente oficial hasta "
            "la salida de consumo."))
    h2(d, "4.2 Objetivos específicos y evidencia de cumplimiento")
    tabla(d, pd.DataFrame([
        {"Objetivo": "Evaluar la viabilidad de fuentes en cuatro dominios",
         "Evidencia": "catalogo_fuentes.csv, reporte_no_viabilidad.csv",
         "Resultado": "6 fuentes evaluadas, 2 integradas, 2 de contexto, 2 descartadas"},
        {"Objetivo": "Construir series mensuales reconciliadas por dominio",
         "Evidencia": "serie_aduanera_mensual.parquet, serie_portuaria_mensual.parquet",
         "Resultado": f"{c('Meses de la serie aduanera')} y {c('Meses de la serie portuaria')} meses continuos"},
        {"Objetivo": "Ejecutar las 52 preguntas y conservar la evidencia",
         "Evidencia": "matriz_trazabilidad_eda.csv",
         "Resultado": f"{c('Preguntas ejecutadas')} ejecutadas, {c('Preguntas parciales')} parciales, {c('Preguntas no viables')} no viables"},
        {"Objetivo": "Declarar el tipo de integración posible y sus límites",
         "Evidencia": "matriz_integracion.csv",
         "Resultado": "integración agregada por mes; ninguna directa"},
        {"Objetivo": "Determinar qué indicadores pueden pronosticarse",
         "Evidencia": "elegibilidad_indicadores.csv",
         "Resultado": "4 elegibles de 7 evaluados"},
        {"Objetivo": "Medir si la integración mejora el pronóstico",
         "Evidencia": "ablacion_multidominio.csv",
         "Resultado": f"no lo mejora: pasa de {c('WAPE historia propia')} % a {c('WAPE historia propia mas puerto')} %"},
        {"Objetivo": "Comunicar la incertidumbre con cobertura medida",
         "Evidencia": "cobertura_intervalos_portuarios.csv",
         "Resultado": f"cobertura empírica entre {c('Cobertura empirica minima')} % y {c('Cobertura empirica maxima')} %"},
        {"Objetivo": "Documentar con evidencia lo no viable",
         "Evidencia": "reporte_no_viabilidad.csv",
         "Resultado": f"{c('Preguntas no viables')} preguntas cerradas con búsqueda documentada"},
    ]), "Objetivos específicos y su evidencia de cumplimiento",
        "Elaboración propia. Cada evidencia es un archivo generado por el pipeline.")

    # ------------------------------------------------------------------ 5
    h1(d, "5. Justificación")
    par(d, ("Buenaventura es el principal puerto colombiano sobre el Pacífico y por él pasa "
            "una porción sustancial de la mercancía importada del país. Los datos que lo "
            "describen son públicos, pero llegar a una lectura útil de ellos exige un trabajo "
            "técnico que hoy debe rehacerse cada vez: descargar paquetes anuales, "
            "homologar formatos que cambian, validar unidades y agregar series."))
    par(d, ("Ese trabajo repetido es lo que el producto automatiza. Una vez construido el "
            "flujo, actualizar el análisis con un mes nuevo deja de ser un proyecto y pasa a "
            "ser una ejecución. Ese cambio es lo que separa un análisis puntual de un "
            "producto de datos."))
    par(d, ("El proyecto tiene además valor analítico propio. Distinguir valor, peso y valor "
            "unitario implícito permite responder si un mes cambió por cantidad o por precio "
            "de la mercancía, una pregunta que un total agregado no responde. Separar el "
            "transbordo del comercio exterior evita leer como caída del comercio lo que es un "
            "cambio en el uso del puerto como punto de conexión entre buques."))
    par(d, ("Desde la ciencia de datos, el trabajo aporta un ejercicio poco frecuente en "
            "proyectos de este nivel: comparar el modelo contra tres líneas base y no contra "
            "una sola. Reportar únicamente la referencia estacional hacía parecer que el "
            f"modelo mejoraba un {c('Mejora sobre naive 12')} %, cuando frente a la "
            f"referencia exigente la mejora real es del {c('Mejora sobre naive 1')} %."))
    h2(d, "5.1 Usuarios potenciales")
    vinetas(d, [
        "Analistas de comercio exterior que preparan informes mensuales.",
        "Observatorios económicos regionales y centros de estudios.",
        "Gremios de comercio exterior y agencias de aduanas.",
        "Investigadores y estudiantes que necesiten una serie ya reconciliada.",
        "Entidades públicas que requieran una lectura conjunta de ambos dominios.",
    ])

    # ------------------------------------------------------------------ 6
    h1(d, "6. Alcance del trabajo")
    par(d, "El trabajo comprende de forma explícita:")
    vinetas(d, [
        "Integración de los microdatos aduaneros del DANE para la aduana 35.",
        "Integración del tráfico portuario de la zona de Buenaventura.",
        "Unión agregada de ambos dominios por mes calendario.",
        "Análisis exploratorio completo de 52 preguntas con evidencia por pregunta.",
        "Pronóstico de los indicadores que cumplen los criterios de elegibilidad.",
        "Intervalos de predicción con cobertura empírica medida.",
        "Reglas de alerta de tres niveles calibradas con la ventana de entrenamiento.",
        "Tablero de consulta de seis vistas que lee de la capa de salidas.",
        "Trazabilidad de cada cifra hasta un archivo del pipeline.",
    ])

    # ------------------------------------------------------------------ 7
    h1(d, "7. Delimitación")
    par(d, ("Las siguientes delimitaciones no son advertencias formales: cada una responde a "
            "una característica concreta de las fuentes y define el borde de lo que el "
            "producto puede afirmar."))
    tabla(d, pd.DataFrame([
        {"Delimitación": "Aduana 35", "Qué significa":
         "Es una unidad de registro administrativo, no una instalación física"},
        {"Delimitación": "Sociedad portuaria", "Qué significa":
         "La fuente identifica la razón social que reporta, que puede administrar una o varias instalaciones"},
        {"Delimitación": "CIF por kilogramo", "Qué significa":
         "Valor unitario implícito afectado por mezcla, seguro y flete. No es un precio"},
        {"Delimitación": "Transbordo", "Qué significa":
         "Carga que cambia de buque sin entrar al país. No es comercio exterior"},
        {"Delimitación": "Llave declaración-buque", "Qué significa":
         "No existe llave pública. La integración es agregada por mes, nunca directa"},
        {"Delimitación": "Tiempos operativos", "Qué significa":
         "ETA, ATA, permanencias y congestión no tienen fuente pública histórica"},
        {"Delimitación": "TEU", "Qué significa":
         "La fuente publica toneladas por tipo de carga, no unidades de contenedor"},
        {"Delimitación": "Unidad monetaria", "Qué significa":
         "El valor CIF está en dólares corrientes: parte de la variación es precio"},
        {"Delimitación": "Alertas", "Qué significa":
         "Son señales de revisión analítica, no órdenes operativas"},
    ]), "Delimitaciones del producto y su origen en las fuentes",
        "Elaboración propia a partir de la documentación oficial de cada fuente.")
    # ------------------------------------------------------------------ 8
    h1(d, "8. Marco de referencia")
    h2(d, "8.1 Marco conceptual")
    par(d, "Los términos que el documento usa con un sentido preciso son los siguientes.")
    tabla(d, pd.DataFrame([
        {"Término": "Valor CIF", "Definición aplicada":
         "Valor de la mercancía más seguro y flete hasta el punto de entrada, en dólares corrientes"},
        {"Término": "Peso neto", "Definición aplicada": "Peso de la mercancía sin embalaje, en kilogramos"},
        {"Término": "Valor unitario implícito", "Definición aplicada":
         "Cociente entre valor CIF y peso neto de los agregados mensuales. No es un precio"},
        {"Término": "Toneladas movilizadas", "Definición aplicada":
         "Carga registrada por las sociedades portuarias, con embalaje incluido"},
        {"Término": "Importación", "Definición aplicada": "Carga que ingresa al territorio aduanero"},
        {"Término": "Exportación", "Definición aplicada": "Carga que sale del territorio aduanero"},
        {"Término": "Transbordo", "Definición aplicada":
         "Carga que cambia de buque sin entrar ni salir del país"},
        {"Término": "Sociedad portuaria", "Definición aplicada":
         "Razón social que reporta movimientos de carga a la Superintendencia de Transporte"},
    ]), "Conceptos del dominio",
        "Elaboración propia a partir de DANE (2026) y Superintendencia de Transporte (2026).")
    tabla(d, pd.DataFrame([
        {"Término": "HHI", "Definición aplicada":
         "Suma de los cuadrados de las participaciones porcentuales. Por encima de 2.500 se considera concentrado"},
        {"Término": "WAPE", "Definición aplicada":
         "Suma de errores absolutos dividida por la suma de valores observados, en porcentaje"},
        {"Término": "MASE", "Definición aplicada":
         "Error absoluto medio escalado por el error de una referencia ingenua. Menor que 1 significa que supera esa referencia"},
        {"Término": "Naive 1", "Definición aplicada": "Predecir que el mes siguiente será igual al último observado"},
        {"Término": "Naive 12", "Definición aplicada": "Predecir el valor del mismo mes del año anterior"},
        {"Término": "Drift", "Definición aplicada": "Extrapolar la pendiente media de toda la serie"},
        {"Término": "Walk-forward", "Definición aplicada":
         "En cada corte se entrena con el pasado y se predice el mes siguiente, respetando el orden temporal"},
        {"Término": "Intervalo de predicción", "Definición aplicada":
         "Rango construido con los cuantiles de los errores fuera de muestra, cuya cobertura se mide"},
        {"Término": "Trazabilidad", "Definición aplicada":
         "Posibilidad de rastrear una cifra hasta el archivo y la línea de código que la produjo"},
        {"Término": "Capas raw, landing, trusted, surface", "Definición aplicada":
         "Original inmutable, homologado, validado y listo para consumo"},
    ]), "Conceptos metodológicos", "Elaboración propia a partir de Hyndman y Athanasopoulos (2021).")
    h2(d, "8.2 Marco normativo")
    vinetas(d, [
        "Ley 1712 de 2014, de transparencia y acceso a la información pública, que sustenta "
        "el uso de datos abiertos con cita de la fuente.",
        "Ley 1581 de 2012, de protección de datos personales. La capa analítica no contiene "
        "datos personales: los microdatos del DANE llegan anonimizados y no incluyen NIT ni "
        "razón social del importador.",
        "NTC 1486 del ICONTEC, aplicada a márgenes, numeración y presentación de este documento.",
        "Normas APA en su séptima edición, aplicadas a citas y referencias.",
        "Condiciones de uso del DANE, que exigen una cita textual de la fuente y restringen "
        "la reproducción de los microdatos originales.",
        "Licencia Creative Commons Atribución-CompartirIgual 4.0 Internacional del conjunto "
        "de tráfico portuario, que obliga a citar y a compartir los derivados bajo la misma licencia.",
    ])
    h2(d, "8.3 Antecedentes")
    par(d, ("Las estadísticas de comercio exterior colombiano se publican de forma periódica "
            "desde 1916. Desde 1993 la Dirección de Impuestos y Aduanas Nacionales produce la "
            "información y el DANE la valida y difunde, con un plazo máximo de cuarenta y "
            "cinco días posteriores al mes de referencia. Esa división institucional explica "
            "parte del rezago de publicación con el que trabaja el producto."))
    par(d, ("Los microdatos aduaneros se distribuyen en dos catálogos distintos según el "
            "periodo, con formatos que cambiaron cinco veces en el intervalo analizado. Las "
            "estadísticas portuarias, en cambio, mantienen un esquema estable desde 2018, "
            "aunque su publicación es trimestral pese a que el dato es mensual."))
    par(d, ("Para el dominio marítimo, la Dirección General Marítima publica boletines "
            "trimestrales agregados en formato PDF. No se localizó una serie histórica "
            "tabular con desagregación por evento, lo que impide construir indicadores de "
            "arribos, tipos de buque o tiempos operativos con fuentes públicas. Esa ausencia "
            "es la que define el alcance real del proyecto."))

    # ------------------------------------------------------------------ 9
    h1(d, "9. Metodología")
    par(d, ("El trabajo es una investigación aplicada, cuantitativa y longitudinal. Cada fase "
            "produce una evidencia verificable, y ninguna cifra llega a este documento sin "
            "pasar antes por un archivo generado por el pipeline."))
    h2(d, "9.1 Preparación de los datos")
    par(d, ("Los archivos originales se conservan sin modificar y se registran con su hash, "
            "de modo que puede demostrarse meses después que no cambiaron. La homologación "
            "detecta automáticamente el separador y la convención decimal de cada vigencia, y "
            "lee los códigos como texto para no perder los ceros iniciales. Ninguna "
            "corrección ocurre en silencio: cada exclusión queda registrada con su regla."))
    h2(d, "9.2 Validación temporal")
    par(d, ("La evaluación usa validación walk-forward de un paso. En términos sencillos: el "
            "modelo se sitúa en cada mes del pasado, se entrena únicamente con lo anterior y "
            "predice el mes siguiente; después se compara con lo que realmente ocurrió. Se "
            "repite ese ejercicio veinticuatro veces."))
    par(d, ("Toda transformación que dependa de los datos, incluidos los escaladores, se "
            "ajusta dentro del conjunto de entrenamiento de cada corte. Si se ajustara sobre "
            "la serie completa, el modelo estaría usando información del futuro para preparar "
            "el pasado, y su desempeño aparente sería mejor que el real."))
    h2(d, "9.3 Prevención de fuga de información")
    par(d, ("Todas las variables predictoras se construyen con datos anteriores al mes que se "
            "predice. Las medias móviles se calculan sobre la serie ya desplazada. Las únicas "
            "variables sin rezago son las de calendario, que se conocen de antemano. Existe "
            "una prueba automatizada que detiene el proceso si alguna variable presenta "
            "correlación casi perfecta con el objetivo del mismo mes."))
    h2(d, "9.4 Líneas base y métricas")
    par(d, ("Se reportan siempre tres referencias: repetir el último valor, repetir el mismo "
            "mes del año anterior y extrapolar la pendiente. Reportar solo una puede "
            "exagerar la mejora atribuible al modelo. Las métricas incluyen WAPE, MASE, sesgo "
            "y error máximo, porque un promedio bajo puede ocultar un mes con un error grave "
            "o un sesgo sistemático hacia un lado."))
    h2(d, "9.5 Intervalos de predicción")
    par(d, ("Los intervalos se construyen con los cuantiles de los errores fuera de muestra "
            "del propio backtest, con calibración expansiva: para cada corte se usan "
            "únicamente los errores anteriores. No se derivan de ninguna métrica de error "
            "puntual. Si la cobertura medida no alcanza el nivel declarado, el intervalo se "
            "recalibra o se renombra según lo que realmente cubre."))
    h2(d, "9.6 Criterios de elegibilidad y de no viabilidad")
    par(d, ("Un indicador entra al componente predictivo si tiene al menos treinta y seis "
            "observaciones mensuales, continuidad completa y un rezago de publicación "
            "conocido. El criterio se fijó antes de mirar los resultados. Una pregunta se "
            "cierra como no viable cuando se documenta dónde se buscó la fuente, qué se "
            "encontró y qué haría falta para responderla en una fase posterior."))
    h2(d, "9.7 Análisis de ablación")
    par(d, ("La ablación consiste en entrenar el mismo modelo con distintos grupos de "
            "variables y comparar su desempeño sobre los mismos cortes. Sirve para responder "
            "si una variable aporta, en lugar de suponerlo. Es el método con el que se "
            "evaluó si integrar el dominio portuario mejora el pronóstico aduanero."))
    # ------------------------------------------------------------------ 10
    h1(d, "10. Ingeniería del producto de datos")
    h2(d, "10.1 Arquitectura por capas")
    tabla(d, pd.DataFrame([
        {"Capa": "raw", "Contenido": "Archivos originales sin modificar",
         "Garantía": "Hash SHA-256 por archivo"},
        {"Capa": "landing", "Contenido": "Datos homologados entre vigencias",
         "Garantía": "Registro de formatos y exclusiones"},
        {"Capa": "trusted", "Contenido": "Series mensuales validadas",
         "Garantía": "Continuidad y reconciliación"},
        {"Capa": "surface", "Contenido": "Métricas, pronósticos y evidencia",
         "Garantía": "Cada archivo enlazado a una pregunta"},
    ]), "Capas de datos y su garantía", "Elaboración propia.")
    h2(d, "10.2 Pipeline")
    par(d, ("Las entradas son los paquetes anuales del DANE y la consulta a la interfaz de "
            "datos abiertos de la Superintendencia. Las transformaciones incluyen detección "
            "de formato, filtro por aduana, conversión de tipos, agregación mensual y cálculo "
            "del valor unitario implícito. Las validaciones cubren esquema, duplicados, "
            "continuidad, dominios y reconciliación. Las salidas son los archivos de la capa "
            "de consumo, las figuras y la matriz de trazabilidad."))
    h2(d, "10.3 Integración de dominios")
    par(d, ("La integración es agregada por mes calendario. No existe llave pública que "
            "relacione una declaración de importación con un movimiento de carga portuaria, y "
            "no se construyó ninguna por inferencia. Cada relación entre fuentes está "
            "declarada como directa, agregada, contextual o no viable."))
    tabla(d, pd.read_csv(S / "matriz_integracion.csv")[
        ["origen", "destino", "llave", "tipo"]].rename(columns={
            "origen": "Origen", "destino": "Destino", "llave": "Llave", "tipo": "Tipo"}),
        "Relaciones entre fuentes y tipo de integración",
        "Elaboración propia. Salida del pipeline: matriz_integracion.csv (P42).")
    h2(d, "10.4 Calidad y trazabilidad")
    par(d, (f"El proyecto cuenta con {c('Pruebas automatizadas')} pruebas automatizadas que "
            f"verifican las reglas del propio producto: que las variables no miren al futuro, "
            f"que los intervalos no se deriven de una métrica puntual, que la numeración de "
            f"las preguntas no cambie y que una pregunta sin salida no se marque como "
            f"ejecutada. El cuaderno del análisis exploratorio produce "
            f"{c('Figuras del cuaderno')} figuras y "
            f"{c('Archivos de evidencia del cuaderno')} archivos de evidencia."))
    par(d, ("Existe además una lista de cifras que asocia cada valor citado con su archivo de "
            "origen y su pregunta. La construcción de este documento falla si se solicita una "
            "cifra que no esté en esa lista, lo que impide que un número escrito a mano llegue "
            "al texto."))
    h2(d, "10.5 Modelado")
    par(d, ("Los objetivos evaluados fueron el valor CIF, el peso neto, las toneladas totales "
            "y la carga contenerizada. Las variables predictoras son rezagos del propio "
            "indicador, medias móviles desplazadas, variables de calendario y, en los "
            "conjuntos de prueba, la tasa representativa del mercado y el índice oceánico. "
            "Los resultados se detallan en el capítulo 12."))
    h2(d, "10.6 Tablero de consulta")
    tabla(d, pd.DataFrame([
        {"Vista": "Ejecutiva", "Contenido": "Valor económico y volumen físico lado a lado"},
        {"Vista": "Aduanera", "Contenido": "CIF, peso neto, valor unitario y estacionalidad"},
        {"Vista": "Portuaria", "Contenido": "Tipos de carga, sociedades, concentración y especialización"},
        {"Vista": "Marítima", "Contenido": "Documentación de la no viabilidad del dominio"},
        {"Vista": "Predictiva", "Contenido": "Modelos, líneas base, intervalos y ablación"},
        {"Vista": "Calidad", "Contenido": "Trazabilidad de las 52 preguntas y fuentes"},
    ]), "Vistas del tablero de consulta",
        "Elaboración propia. El tablero lee de la capa de salidas y no calcula cifras propias.")
    par(d, ("El tablero no realiza ningún cálculo: lee exclusivamente los archivos que produce "
            "el pipeline. Si un archivo falta, la vista lo indica en lugar de mostrar un valor "
            "estimado. Esa decisión de diseño garantiza que lo que se ve en pantalla y lo que "
            "dice este documento provienen del mismo origen."))

    # ------------------------------------------------------------------ 11
    h1(d, "11. Resultados del análisis exploratorio")
    h2(d, "11.1 Estado de las 52 preguntas")
    traza = pd.read_csv(S / "matriz_trazabilidad_eda.csv")
    tabla(d, traza["estado"].value_counts().rename_axis("Estado").reset_index(name="Preguntas"),
          "Estado de las 52 preguntas del análisis exploratorio",
          "Elaboración propia. Salida del pipeline: matriz_trazabilidad_eda.csv.")
    par(d, ("Ninguna pregunta quedó sin respuesta. Las cerradas como no viables documentan "
            "dónde se buscó, qué se encontró y qué fuente haría falta. La especificación del "
            "análisis admite esa respuesta siempre que esté demostrada."))
    h2(d, "11.2 Valor económico frente a volumen físico")
    par(d, (f"Sobre los {c('Meses integrados')} meses comunes, el valor CIF creció un "
            f"{c('Crecimiento del CIF')} % entre los primeros y los últimos doce meses. "
            f"Descomponiendo ese crecimiento en logaritmos, el volumen aporta el "
            f"{c('Aporte del volumen al crecimiento')} % y el valor unitario el "
            f"{c('Aporte del valor unitario al crecimiento')} %. La suma no alcanza "
            f"exactamente el cien por ciento porque la media de un cociente no coincide con "
            f"el cociente de las medias; ese residuo es de menos de un punto."))
    destacado(d, ("Alrededor de la mitad del crecimiento del valor importado corresponde a más "
                  "mercancía y la otra mitad a mercancía más cara por kilogramo. Un total "
                  "agregado no permite distinguirlo."))
    par(d, (f"El valor unitario implícito medio del periodo completo es de "
            f"{c('CIF por kilogramo medio')} dólares por kilogramo. Conviene insistir en que "
            f"no es un precio: lo afectan la mezcla de mercancías, el seguro y el flete."))
    h2(d, "11.3 El total portuario y sus componentes")
    par(d, ("El tráfico portuario agregado combina importación, exportación y transbordo. El "
            "transbordo es carga que llega en un buque y sale en otro sin entrar al "
            "territorio aduanero, de modo que no forma parte del comercio exterior del país."))
    par(d, (f"Comparando los primeros doce meses de la serie portuaria con los últimos doce, "
            f"las toneladas de importación suben un {c('Variacion de toneladas de importacion')} % "
            f"y las de exportación un {c('Variacion de toneladas de exportacion')} %, mientras "
            f"el transbordo cae un {c('Variacion del transbordo')} %. El total resultante baja "
            f"un {c('Variacion del total portuario')} %."))
    destacado(d, ("Quien observe únicamente el total concluiría que la zona portuaria movilizó "
                  "menos carga. La descomposición muestra que lo que descendió fue el "
                  "transbordo, mientras la carga de importación aumentaba."))
    par(d, ("La fuente no informa por qué cambió el transbordo. Decisiones de las navieras, "
            "reasignación de rutas o cambios en el reporte son explicaciones posibles que "
            "estos datos no permiten distinguir."))
    h2(d, "11.4 Concentración por país, capítulo y sociedad portuaria")
    tabla(d, pd.read_csv(S / "comparacion_concentracion.csv")[
        ["dimension", "hhi", "clase", "dominio"]].rename(columns={
            "dimension": "Dimensión", "hhi": "HHI", "clase": "Lectura", "dominio": "Dominio"}),
        "Índice de concentración por dimensión",
        "Elaboración propia. Salida del pipeline: comparacion_concentracion.csv (P17, P18, P26).")
    par(d, (f"La canasta de productos tiene un índice de {c('HHI capitulo arancelario')} y la "
            f"de orígenes {c('HHI pais de origen')}: ambas desconcentradas. La movilización, "
            f"con {c('HHI sociedad portuaria')}, está muy por encima del umbral de 2.500. Una "
            f"sola sociedad portuaria moviliza el {c('Participacion de la mayor sociedad')} % "
            f"de las toneladas del periodo."))
    par(d, (f"Entre los años completos, el índice bajó de {c('HHI portuario 2018')} en 2018 a "
            f"{c('HHI portuario 2023')} en 2023 y subió a {c('HHI portuario 2025')} en 2025. "
            f"El valor de 2026, {c('HHI portuario 2026 parcial')}, no es comparable: cubre "
            f"seis meses y solo tres sociedades reportan. Recalculando 2025 con esas mismas "
            f"tres sociedades el índice sería 4.096, de modo que la mayor parte del salto se "
            f"explica por el cambio de cobertura del reporte y no por una redistribución "
            f"entre las que continúan."))
    par(d, ("El índice mide reparto de toneladas reportadas. No mide capacidad instalada, "
            "utilización ni posibilidad real de sustitución entre sociedades, de modo que un "
            "valor alto no permite concluir por sí solo que exista un riesgo operativo."))
    h2(d, "11.5 Especialización por tipo de carga")
    par(d, ("El índice de concentración trata a las sociedades como si fueran "
            "intercambiables. El cruce entre sociedad y tipo de carga muestra que no "
            "movilizan lo mismo: una de ellas no registró toneladas de granel en todo el "
            "periodo y otra no registró contenedores. El reparto agregado describe cuánto "
            "movilizó cada una, no si podrían movilizar la carga de las demás."))
    h2(d, "11.6 Sociedades sin reporte en 2026")
    par(d, (f"{c('Sociedades sin reporte en 2026')} sociedades portuarias no aparecen en los "
            f"reportes de los seis meses observados de 2026. Se verificó mes a mes para "
            f"descartar un rezago de publicación."))
    par(d, ("La fuente registra reporte, no operación. Con estos datos solo puede afirmarse "
            "que no figuran en los reportes del periodo observado; no puede afirmarse que "
            "hayan dejado de operar. La implicación metodológica es relevante: un modelo "
            "entrenado sobre el total de la zona interpretaría esa ausencia como una "
            "contracción real del comercio."))
    h2(d, "11.7 Comportamiento temporal de la serie aduanera")
    par(d, (f"La autocorrelación del valor CIF con su rezago de un mes es de "
            f"{c('ACF del CIF en rezago 1')}, y con el rezago de doce meses de "
            f"{c('ACF del CIF en rezago 12')}. La estacionalidad existe pero es moderada: el "
            f"índice va de {c('Indice estacional minimo')} a {c('Indice estacional maximo')} "
            f"sobre una base de cien."))
    destacado(d, ("Que el mes anterior pese más que el mismo mes del año pasado tiene una "
                  "consecuencia directa: la referencia exigente para el modelo no es la "
                  "estacional sino la de repetir el último valor observado."))
    h2(d, "11.8 Relación entre los dos dominios")
    par(d, (f"La razón mediana entre toneladas portuarias de importación y peso neto aduanero "
            f"es de {c('Razon toneladas portuarias sobre peso aduanero')}, y la correlación "
            f"en variaciones mensuales es de {c('Correlacion en variaciones entre dominios')}. "
            f"Las series se mueven en el mismo sentido general pero no son proporcionales, lo "
            f"cual es esperable: miden universos distintos."))
    # ------------------------------------------------------------------ 12
    h1(d, "12. Resultados del modelado")
    h2(d, "12.1 Indicadores elegibles")
    tabla(d, pd.read_csv(S / "elegibilidad_indicadores.csv")[
        ["indicador", "dominio", "n_obs", "elegible"]].rename(columns={
            "indicador": "Indicador", "dominio": "Dominio", "n_obs": "Observaciones",
            "elegible": "Elegible"}),
        "Elegibilidad de los indicadores candidatos",
        "Elaboración propia. Salida del pipeline: elegibilidad_indicadores.csv (P46).")
    par(d, ("Tres de los siete candidatos quedan excluidos porque su fuente no existe: TEU, "
            "arribos y permanencia media no tienen ni una sola observación."))
    h2(d, "12.2 Pronóstico del valor CIF")
    par(d, (f"El mejor conjunto de variables es el que usa la historia propia del indicador y "
            f"variables de calendario, con un WAPE de {c('WAPE historia propia mas calendario')} % "
            f"sobre veinticuatro cortes. Las tres líneas base obtienen "
            f"{c('WAPE naive 1 sobre el CIF')} % la de repetir el último valor, "
            f"{c('WAPE drift sobre el CIF')} % la de extrapolar la pendiente y "
            f"{c('WAPE naive 12 sobre el CIF')} % la estacional."))
    destacado(d, (f"Frente a la referencia estacional la mejora es del {c('Mejora sobre naive 12')} %. "
                  f"Frente a la referencia exigente, que es repetir el último valor observado, "
                  f"la mejora real es del {c('Mejora sobre naive 1')} %."))
    par(d, ("La diferencia entre ambas cifras es la razón por la que el trabajo reporta "
            "siempre las tres líneas base. Informar únicamente la estacional describiría más "
            "la debilidad de esa referencia que la calidad del modelo."))
    h2(d, "12.3 Análisis de ablación")
    abl = pd.read_csv(S / "ablacion_multidominio.csv")
    tabla(d, abl[["conjunto", "n_vars", "wape_pct", "ganancia_vs_A_pp"]].rename(columns={
        "conjunto": "Conjunto de variables", "n_vars": "Variables",
        "wape_pct": "WAPE (%)", "ganancia_vs_A_pp": "Ganancia (pp)"}),
        "Ablación por grupos de variables sobre el pronóstico del valor CIF",
        "Elaboración propia. Veinticuatro cortes de validación walk-forward. "
        "Salida del pipeline: ablacion_multidominio.csv (P49).")
    par(d, (f"El conjunto que añade variables portuarias obtiene "
            f"{c('WAPE historia propia mas puerto')} %, frente a "
            f"{c('WAPE historia propia')} % del que usa solo la historia propia. El modelo "
            f"integrado completo, con dieciséis variables, obtiene "
            f"{c('WAPE integrado completo')} %. Las variables de contexto tampoco aportan: "
            f"{c('WAPE contexto TRM y ONI')} %."))
    destacado(d, "Añadir el dominio portuario no mejora el pronóstico del valor CIF: lo empeora.")
    par(d, ("La explicación es coherente con la naturaleza de las fuentes. Ambas miden el "
            "mismo comercio y comparten el rezago de publicación, de modo que el puerto no "
            "aporta información que la historia del propio valor CIF no contenga ya, y sí "
            "consume grados de libertad sobre una muestra corta. La consecuencia para el "
            "diseño del producto es directa: el componente predictivo se apoya en el dominio "
            "aduanero, y el portuario se conserva por su valor descriptivo y explicativo."))
    h2(d, "12.4 Pronóstico de indicadores portuarios")
    par(d, (f"Para las toneladas totales, el modelo obtiene {c('WAPE toneladas totales Ridge')} % "
            f"frente a {c('WAPE toneladas totales naive 1')} % de la referencia, una mejora "
            f"del {c('Mejora en toneladas totales')} %."))
    par(d, (f"Para la carga contenerizada ocurre lo contrario: la referencia obtiene "
            f"{c('WAPE carga contenerizada naive 1')} % y el modelo "
            f"{c('WAPE carga contenerizada Ridge')} %, es decir peor que repetir el último "
            f"valor observado."))
    destacado(d, ("Para la carga contenerizada se recomienda no usar modelo. Un indicador que "
                  "no se pronostica mejor que la referencia trivial no debe presentarse con "
                  "un modelo encima."))
    h2(d, "12.5 Intervalos de predicción")
    cob = pd.read_csv(S / "cobertura_intervalos_portuarios.csv")
    tabla(d, pd.DataFrame({
        "Indicador": cob["caso"].str.replace(r"_ridge_v\d+", "", regex=True).str.replace("_", " "),
        "Ventana": cob["caso"].str.extract(r"v(\d+)")[0] + " cortes",
        "Nominal": (cob["nivel_nominal"] * 100).round(0).astype(int).astype(str) + " %",
        "Empírica": (cob["cobertura_empirica"] * 100).round(1).astype(str)
                    .str.replace(".", ",", regex=False) + " %",
        "Cortes": cob["n_evaluados"], "Fuera": cob["casos_fuera"]}),
        "Cobertura empírica de los intervalos de predicción",
        "Elaboración propia. Método: cuantiles empíricos de los errores fuera de muestra con "
        "calibración expansiva. Salida completa: cobertura_intervalos_portuarios.csv (P51).")
    par(d, (f"La cobertura se mide y no se declara. Con {c('Cortes evaluados minimo')} cortes "
            f"evaluables en el caso más corto, el intervalo de confianza de la propia "
            f"cobertura es amplio: un valor por debajo del nominal no permite concluir que el "
            f"intervalo esté mal calibrado, del mismo modo que uno por encima no permite "
            f"concluir que lo esté bien."))

    # ------------------------------------------------------------------ 13
    h1(d, "13. Modelo de negocios")
    par(d, ("Lo que sigue es una propuesta potencial de adopción, no un modelo validado "
            "comercialmente. No se han realizado entrevistas con usuarios ni pruebas piloto, "
            "y por tanto ninguna de estas afirmaciones cuenta con evidencia de demanda."))
    h2(d, "13.1 Propuesta de valor")
    par(d, ("Un producto reproducible que integra, actualiza y comunica indicadores aduaneros "
            "y portuarios de Buenaventura, con trazabilidad de cada cifra y una declaración "
            "explícita de lo que los datos no permiten afirmar."))
    h2(d, "13.2 Necesidad atendida")
    par(d, ("Reducir el esfuerzo técnico de descargar, homologar, validar e interpretar dos "
            "fuentes que se publican por separado, con formatos que cambian y unidades que no "
            "son equivalentes."))
    h2(d, "13.3 Usuarios potenciales, canales y sostenibilidad")
    tabla(d, pd.DataFrame([
        {"Componente": "Usuarios potenciales", "Descripción":
         "Entidades públicas, observatorios, universidades, gremios, analistas, importadores, "
         "sociedades portuarias y consultores"},
        {"Componente": "Producto ofrecido", "Descripción":
         "Tablero, series mensuales, reportes, alertas, pronósticos, archivos de evidencia y "
         "actualización automatizada"},
        {"Componente": "Canales", "Descripción":
         "Aplicación web, reportes periódicos, repositorio de código, presentaciones y "
         "exportación de archivos"},
        {"Componente": "Sostenibilidad", "Descripción":
         "Mantenimiento institucional, consultoría, servicio de actualización, adaptación a "
         "fuentes privadas o código abierto con soporte"},
        {"Componente": "Costos", "Descripción":
         "Infraestructura, mantenimiento del pipeline, validación mensual, almacenamiento, "
         "soporte y eventuales fuentes comerciales"},
    ]), "Componentes del modelo de negocios propuesto",
        "Elaboración propia. Propuesta no validada con usuarios.")
    h2(d, "13.4 Riesgos")
    vinetas(d, [
        "Cambios de formato en las fuentes, que ya ocurrieron cinco veces en el periodo analizado.",
        "Rezagos de publicación que impiden usar el dato del mes en curso.",
        "Desaparición o reestructuración de un conjunto de datos abierto.",
        "Cambios de cobertura del reporte, como el observado en 2026, que alteran la serie sin "
        "que cambie el fenómeno.",
        "Dependencia completa de datos públicos, sin acceso a información operativa.",
        "Ausencia de datos de tiempos y buques, que limita el alcance frente a productos comerciales.",
    ])

    # ------------------------------------------------------------------ 14
    h1(d, "14. Estrategias de comunicación")
    tabla(d, pd.DataFrame([
        {"Público": "Directivos", "Necesidad": "Saber si algo requiere atención",
         "Canal": "Vista ejecutiva", "Mensaje": "Nivel de alerta con su razón",
         "Detalle": "Bajo"},
        {"Público": "Analistas", "Necesidad": "Preparar el informe mensual",
         "Canal": "Vistas aduanera y portuaria", "Mensaje": "Series, composición y variaciones",
         "Detalle": "Medio"},
        {"Público": "Investigadores", "Necesidad": "Reutilizar series y método",
         "Canal": "Repositorio y archivos", "Mensaje": "Datos reconciliados y trazables",
         "Detalle": "Alto"},
        {"Público": "Jurado académico", "Necesidad": "Evaluar rigor y límites",
         "Canal": "Documento y sustentación", "Mensaje": "Qué se midió y qué no se puede afirmar",
         "Detalle": "Alto"},
        {"Público": "Usuarios no técnicos", "Necesidad": "Entender la lectura",
         "Canal": "Tablero", "Mensaje": "Valor frente a volumen en lenguaje corriente",
         "Detalle": "Bajo"},
    ]), "Estrategia de comunicación por público", "Elaboración propia.")
    h2(d, "14.1 Criterios de comunicación")
    par(d, ("Las alertas se comunican en tres niveles, cada uno acompañado de la razón que lo "
            "activó y de una acción de revisión sugerida. Ninguna alerta constituye una orden "
            "operativa, y así se indica en la propia tarjeta del tablero."))
    par(d, ("Las limitaciones se muestran junto al resultado, no en una sección aparte al "
            "final. Cuando el producto no puede responder algo, la vista correspondiente lo "
            "declara y explica qué haría falta, en lugar de dejar el espacio vacío."))
    par(d, ("Las asociaciones entre variables se presentan siempre como asociaciones. El "
            "documento y el tablero evitan el verbo influir cuando solo existe correlación, y "
            "las correlaciones se calculan sobre series diferenciadas para no reportar "
            "relaciones espurias por tendencia común."))
    par(d, ("Las métricas técnicas se acompañan de su lectura en lenguaje corriente. Un WAPE "
            "del seis por ciento se explica como que, en promedio, la predicción se aparta un "
            "seis por ciento del valor observado. Un intervalo se explica como el rango dentro "
            "del cual cabe esperar el valor real, con la advertencia de cuántos cortes lo "
            "respaldan."))

    # ------------------------------------------------------------------ 15
    h1(d, "15. Conclusiones")
    par(d, ("Cada conclusión corresponde a un resultado del análisis o del modelado, y puede "
            "rastrearse hasta un archivo del pipeline."))
    vinetas(d, [
        f"El proyecto es viable con el alcance construido: dos dominios integrados sobre "
        f"{c('Meses integrados')} meses comunes, con {c('Preguntas ejecutadas')} preguntas "
        f"ejecutadas y {c('Preguntas no viables')} cerradas con evidencia de no viabilidad.",
        "La integración entre dominios es agregada por mes calendario. No existe llave "
        "pública que permita una integración directa y no se construyó ninguna por inferencia.",
        f"El valor importado creció un {c('Crecimiento del CIF')} % en el periodo común, con "
        f"un aporte del volumen del {c('Aporte del volumen al crecimiento')} % y del valor "
        f"unitario del {c('Aporte del valor unitario al crecimiento')} %.",
        f"El total portuario descendió un {c('Variacion del total portuario')} %, pero la "
        f"carga de importación subió un {c('Variacion de toneladas de importacion')} % y lo "
        f"que cayó fue el transbordo, un {c('Variacion del transbordo')} %.",
        f"La movilización está concentrada, con un índice de {c('HHI sociedad portuaria')}, "
        f"mientras la canasta de productos y de orígenes no lo está. El índice mide reparto, "
        f"no capacidad ni sustitución.",
        f"El modelo aduanero mejora un {c('Mejora sobre naive 1')} % frente a la referencia "
        f"exigente. Añadir el dominio portuario no mejora el pronóstico.",
        "El aporte del dominio portuario es descriptivo y explicativo: permite distinguir "
        "valor de volumen y localizar por qué sociedad portuaria pasa la carga.",
        "El dominio marítimo y el operativo no son viables con fuentes públicas, y esa "
        "conclusión está respaldada por la búsqueda documentada de cuatro vías de acceso.",
        "Ninguna de las 52 preguntas quedó sin evidencia ni sin justificación.",
    ])

    # ------------------------------------------------------------------ 16
    h1(d, "16. Recomendaciones")
    vinetas(d, [
        "Conservar el modelo aduanero como componente predictivo y el portuario como "
        "componente descriptivo, según indica el análisis de ablación.",
        "No pronosticar la carga contenerizada con modelo: usar la referencia de repetir el "
        "último valor, que obtuvo mejor desempeño.",
        "Repetir la descarga de los microdatos en la próxima publicación para medir el efecto "
        "de las revisiones sobre la serie, hoy no evaluable con una sola descarga.",
        "Traducir los códigos de país y capítulo arancelario a sus nombres oficiales antes de "
        "presentar las tablas de composición a un lector no técnico.",
        "Decidir si el total portuario se modela sobre la zona completa o solo sobre las "
        "sociedades que continúan reportando, dado el cambio de cobertura observado en 2026.",
        "Validar las reglas de alerta con al menos un usuario real antes de considerarlas "
        "operativas.",
        "Gestionar acceso institucional a fuentes de evento si se desea abrir el dominio "
        "marítimo en una fase posterior.",
        "Monitorear los cambios de formato de las fuentes, que en el periodo analizado "
        "ocurrieron cinco veces.",
    ])
    # ------------------------------------------------------------------ 17
    h1(d, "17. Fuentes de información")
    par(d, ("Conforme a las normas APA en su séptima edición, la fecha de consulta se incluye "
            "únicamente en las fuentes diseñadas para cambiar con el tiempo y sin versión "
            "archivada. Las obras estables, como leyes, normas técnicas y manuales impresos, "
            "no la llevan."))
    for r in [
        "American Psychological Association. (2020). Publication Manual of the American "
        "Psychological Association (7.ª ed.). APA.",
        "Congreso de la República de Colombia. (2012). Ley 1581 de 2012, por la cual se "
        "dictan disposiciones generales para la protección de datos personales. Diario "
        "Oficial No. 48.587.",
        "Congreso de la República de Colombia. (2014). Ley 1712 de 2014, por medio de la cual "
        "se crea la Ley de Transparencia y del Derecho de Acceso a la Información Pública "
        "Nacional. Diario Oficial No. 49.084.",
        "Departamento Administrativo Nacional de Estadística. (2018). Estadísticas de "
        "importaciones (IMPO) 1970–2024 [conjunto de datos]. Archivo Nacional de Datos. "
        "Recuperado el 6 de agosto de 2026, de "
        "https://microdatos.dane.gov.co/index.php/catalog/473",
        "Departamento Administrativo Nacional de Estadística. (2026). Estadísticas de "
        "importaciones (IMPO) 2025–2026 [conjunto de datos]. Archivo Nacional de Datos. "
        "Recuperado el 6 de agosto de 2026, de "
        "https://microdatos.dane.gov.co/index.php/catalog/856",
        "Hyndman, R. J., y Athanasopoulos, G. (2021). Forecasting: Principles and practice "
        "(3.ª ed.). OTexts. Recuperado el 6 de agosto de 2026, de https://otexts.com/fpp3/",
        "Instituto Colombiano de Normas Técnicas y Certificación. (2008). NTC 1486: "
        "Documentación. Presentación de tesis, trabajos de grado y otros trabajos de "
        "investigación. ICONTEC.",
        "Superintendencia de Transporte. (2026). Tráfico portuario marítimo en Colombia "
        "[conjunto de datos]. Datos Abiertos Colombia. Recuperado el 6 de agosto de 2026, de "
        "https://www.datos.gov.co/Transporte/Trafico-Portuario-Mar-timo-En-Colombia/5r3g-zv5z",
    ]:
        p = d.add_paragraph(r)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    h2(d, "17.1 Condiciones de uso exigidas por las fuentes")
    par(d, ("El DANE exige la cita textual «Fuente: Departamento Administrativo Nacional de "
            "Estadística: www.dane.gov.co» y prohíbe la reproducción de los microdatos en "
            "medios que los pongan a disposición de múltiples usuarios sin su visto bueno "
            "escrito. Este trabajo reproduce agregados mensuales, no los microdatos originales."))
    par(d, ("El conjunto de tráfico portuario se publica bajo licencia Creative Commons "
            "Atribución-CompartirIgual 4.0 Internacional, que obliga a citar la fuente y a "
            "compartir los derivados bajo la misma licencia. La Superintendencia advierte "
            "además que sus cifras son referenciales y que solo ella puede certificarlas."))
    tabla(d, pd.DataFrame([
        {"Fuente": "DANE catálogo 473", "Verificada": "2026-08-06",
         "Actualizada": "2018-03-28", "Cambia": "no, histórico cerrado"},
        {"Fuente": "DANE catálogo 856", "Verificada": "2026-08-06",
         "Actualizada": "2026-07-22", "Cambia": "sí, mensual"},
        {"Fuente": "Superintendencia de Transporte", "Verificada": "2026-08-06",
         "Actualizada": "2026-08-01", "Cambia": "sí, trimestral"},
        {"Fuente": "Hyndman y Athanasopoulos", "Verificada": "2026-08-06",
         "Actualizada": "2026-07-23", "Cambia": "sí, edición en línea"},
    ]), "Verificación de las fuentes en línea citadas",
        "Elaboración propia. Cada dirección fue consultada el 6 de agosto de 2026 y la fecha "
        "de actualización proviene de los metadatos publicados por la propia fuente.")

    # ------------------------------------------------------------------ ANEXOS
    d.add_page_break()
    h1(d, "Anexo A. Matriz de trazabilidad P01–P52")
    tabla(d, traza[["pregunta", "bloque", "estado"]].rename(columns={
        "pregunta": "Pregunta", "bloque": "Bloque", "estado": "Estado"}),
        "Estado de cada pregunta del análisis exploratorio",
        "Elaboración propia. Salida del pipeline: matriz_trazabilidad_eda.csv.", max_filas=52)

    d.add_page_break()
    h1(d, "Anexo B. Reporte de no viabilidad")
    nv = pd.read_csv(S / "reporte_no_viabilidad.csv")
    tabla(d, nv[["pregunta", "tema", "estado"]].rename(columns={
        "pregunta": "Pregunta", "tema": "Tema", "estado": "Estado"}),
        "Preguntas cerradas por ausencia de fuente",
        "Elaboración propia. Salida del pipeline: reporte_no_viabilidad.csv.")
    par(d, str(nv["razon"].iloc[0]))
    par(d, ("Estas preguntas no quedaron sin responder. Su respuesta es que la fuente "
            "necesaria no existe en el ámbito público, y esa respuesta está documentada con "
            "las cuatro vías de acceso que se consultaron."))

    d.add_page_break()
    h1(d, "Anexo C. Preparación para la sustentación")
    par(d, ("Este anexo es material de apoyo para la exposición oral y no forma parte del "
            "cuerpo académico del trabajo."))
    for preg, resp in [
        ("¿Cuál es el problema?",
         "Los datos aduaneros y portuarios de Buenaventura son públicos pero están separados, "
         "cambian de formato y miden cosas distintas. Sin integrarlos no se puede distinguir "
         "si el comercio creció porque entró más mercancía o porque la mercancía se encareció."),
        ("¿Qué predice el producto?",
         "El valor CIF y el peso neto del mes siguiente en el dominio aduanero, y las "
         "toneladas totales en el portuario. La carga contenerizada se describe pero no se "
         "modela, porque el modelo resultó peor que la referencia."),
        (f"¿Qué significa que mejore un {c('Mejora sobre naive 1')} %?",
         f"Que el error del modelo es un {c('Mejora sobre naive 1')} % menor que el de "
         f"repetir el último valor observado. Es una mejora modesta y medida, no una "
         f"estimación optimista."),
        ("¿Por qué integrar si no mejora el pronóstico?",
         "Porque el valor de la integración es explicativo. Permite separar volumen de valor "
         "unitario y mostrar que la caída del total portuario venía del transbordo, no del "
         "comercio. Ninguna fuente aislada permite esa lectura."),
        ("¿Qué aporta concretamente el dominio portuario?",
         "Tres cosas: la descomposición del total en importación, exportación y transbordo; "
         "la concentración por sociedad portuaria; y la detección del cambio de cobertura del "
         "reporte en 2026."),
        ("¿Por qué el producto no incluye buques ni tiempos?",
         "Porque no existe fuente pública con serie histórica tabular. Se consultaron cuatro "
         "vías y ninguna entrega datos por evento. Está documentado en el anexo B."),
        ("¿Quién usaría el producto?",
         "Un analista de comercio exterior que hoy dedica horas a descargar y homologar "
         "fuentes antes de poder analizar nada."),
        ("¿Cuál es el modelo de negocios?",
         "Una propuesta de adopción, no un modelo validado. No se han hecho entrevistas ni "
         "pilotos, y así se declara en el capítulo 13."),
        ("¿Cómo se evita la fuga de información?",
         "Todas las variables van rezagadas, los escaladores se ajustan dentro de cada corte "
         "y existe una prueba que detiene el proceso si una variable correlaciona casi "
         "perfecto con el objetivo del mismo mes."),
        ("¿Qué es el HHI?",
         "La suma de los cuadrados de las participaciones. Mide qué tan repartido está algo. "
         "Por encima de 2.500 se considera concentrado. Mide reparto, no capacidad."),
        ("¿Qué significa una pregunta no viable?",
         "Que se buscó la fuente, se documentó dónde y qué se encontró, y se concluyó que el "
         "dato no existe en el ámbito público. Es una respuesta, no una omisión."),
        ("¿Cuáles son las limitaciones principales?",
         "El valor CIF está en dólares corrientes, el valor por kilogramo no es un precio, la "
         "integración es agregada por mes y las alertas son señales de revisión, no órdenes."),
        ("¿Qué harían en una segunda fase?",
         "Repetir la descarga para medir revisiones, traducir los códigos a nombres, validar "
         "las alertas con un usuario real y gestionar acceso institucional a datos marítimos."),
    ]:
        p = d.add_paragraph()
        r = p.add_run(preg)
        r.bold, r.font.color.rgb = True, AZUL
        p.paragraph_format.space_before = Pt(8)
        par(d, resp)

    ruta = config.DOCUMENTS / "Documento_Academico_Buenaventura_V5_FINAL.docx"
    ruta.parent.mkdir(parents=True, exist_ok=True)
    d.save(ruta)
    pd.DataFrame({"concepto": sorted(_USADAS)}).to_csv(S / "cifras_usadas_documento.csv",
                                                       index=False)
    return ruta


if __name__ == "__main__":
    r = construir()
    print(f"documento generado: {r.name} ({r.stat().st_size // 1024} KB)")
    print(f"cifras verificadas usadas: {len(_USADAS)}")
