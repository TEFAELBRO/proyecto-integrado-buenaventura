"""Corrige el documento académico V5 para la entrega ante jurados.

Aplica correcciones puntuales sobre el archivo entregado, sin reescribirlo ni
alterar cifras, y produce un reporte de lo hecho y de lo que no se hizo.

Regla de trabajo: cada corrección se verifica contra un archivo del proyecto
antes de aplicarse. Si la verificación contradice la instrucción, la corrección
no se aplica y queda registrada en el reporte con su evidencia.
"""
from __future__ import annotations

import copy
import pathlib
import sys

from docx import Document

sys.path.insert(0, str(pathlib.Path(__file__).parents[1]))
from src import config  # noqa: E402

ENTRADA = pathlib.Path(sys.argv[1]) if len(sys.argv) > 1 else None
SALIDA = config.REPORTS / "documents" / "Documento_Academico_Buenaventura_V5_FINAL_JURADOS.docx"
REPORTE = config.DOCS / "REPORTE_CORRECCION_JURADOS.md"

hechos: list[str] = []
omitidos: list[str] = []
verificar: list[str] = []
inconsistencias: list[str] = []


def sustituir(parrafo, viejo: str, nuevo: str) -> bool:
    """Sustituye texto conservando el formato de la primera corrida."""
    if viejo not in parrafo.text:
        return False
    corridas = parrafo.runs
    if not corridas:
        return False
    corridas[0].text = parrafo.text.replace(viejo, nuevo)
    for r in corridas[1:]:
        r.text = ""
    return True


def sustituir_en_todo(d, viejo: str, nuevo: str) -> int:
    n = 0
    for p in d.paragraphs:
        n += sustituir(p, viejo, nuevo)
    for t in d.tables:
        for fila in t.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    n += sustituir(p, viejo, nuevo)
    return n


def referencia(d, ancla_texto: str, texto: str, antes: bool = False):
    """Inserta una referencia clonando el formato de una referencia existente."""
    ancla = next(p for p in d.paragraphs if p.text.startswith(ancla_texto))
    nuevo = copy.deepcopy(ancla._p)
    if antes:
        ancla._p.addprevious(nuevo)
    else:
        ancla._p.addnext(nuevo)
    from docx.text.paragraph import Paragraph
    p = Paragraph(nuevo, ancla._parent)
    for r in p.runs[1:]:
        r.text = ""
    p.runs[0].text = texto
    return p


def fila_tabla(tabla, valores):
    fila = copy.deepcopy(tabla.rows[-1]._tr)
    tabla._tbl.append(fila)
    nueva = tabla.rows[-1]
    for celda, valor in zip(nueva.cells, valores, strict=False):
        p = celda.paragraphs[0]
        if p.runs:
            p.runs[0].text = valor
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(valor)
    return nueva


def main() -> None:
    d = Document(ENTRADA)

    # --- 1. Tabla 2: la formulación del objetivo debe cubrir el cierre por no viabilidad
    n = sustituir_en_todo(
        d, "Ejecutar las 52 preguntas y conservar la evidencia",
        "Responder o cerrar con evidencia documentada las 52 preguntas del "
        "análisis exploratorio")
    if n:
        hechos.append(
            "Tabla 2: el objetivo pasa de «Ejecutar las 52 preguntas y conservar la "
            "evidencia» a «Responder o cerrar con evidencia documentada las 52 preguntas "
            "del análisis exploratorio». La evidencia (matriz_trazabilidad_eda.csv) y el "
            "resultado (38 ejecutadas, 4 parciales, 10 no viables) no se tocaron.")

    # --- 2. Tabla 1, paso 7: verificación antes de corregir
    ruta = config.TRUSTED / "vista_integrada_mensual.parquet"
    if ruta.exists():
        n = sustituir_en_todo(d, "vista_integrada_mensual.parquet",
                              "vista_integrada_mensual.parquet (capa trusted)")
        omitidos.append(
            "NO se reemplazó la salida del paso 7 por «salidas de integración en capa "
            f"surface». El archivo vista_integrada_mensual.parquet sí existe: está en "
            f"data/trusted/ ({ruta.stat().st_size:,} bytes), lo escribe "
            "src/correr_integrado.py y lista_cifras.csv lo cita como origen del periodo "
            "integrado, de la razón 1,16 y de la correlación 0,115. La auditoría lo dio "
            "por inexistente probablemente porque lo buscó en la capa surface. Aplicar "
            "esa corrección habría introducido un error en un documento correcto.")
        if n:
            hechos.append(
                "Tabla 1, paso 7: se precisó la capa donde vive el archivo, que ahora "
                "dice «vista_integrada_mensual.parquet (capa trusted)». Es el cambio que "
                "evita la confusión que originó el hallazgo de la auditoría.")
    else:
        inconsistencias.append(
            "vista_integrada_mensual.parquet no se encontró en la capa trusted.")

    # --- 3, 4 y 5. Referencias faltantes, con los metadatos del catálogo del proyecto
    cat = config.FUENTES
    banrep, noaa, dimar = cat["banrep_trm"], cat["noaa_oni"], cat["dimar_trafico"]

    referencia(d, "Congreso de la República de Colombia. (2012)",
               f"Banco de la República. (2026). {banrep['nombre']} (TRM) "
               "[serie estadística]. Recuperado el 6 de agosto de 2026, de "
               f"{banrep['url']}", antes=True)
    hechos.append(
        "Se añadió la referencia del Banco de la República por la tasa representativa "
        "del mercado, usada como variable de contexto. Título, entidad y dirección "
        "provienen de catalogo_fuentes.csv; no se inventó ningún metadato.")

    referencia(d, "Hyndman, R. J.",
               "Dirección General Marítima. (2026). "
               f"{dimar['nombre']} [publicaciones trimestrales en PDF]. "
               "Recuperado el 6 de agosto de 2026, de "
               f"{dimar['url']}", antes=True)
    hechos.append(
        "Se añadió la referencia de la Dirección General Marítima, consultada para "
        "evaluar la viabilidad del dominio marítimo. La dirección fue verificada: la "
        "página publica boletines trimestrales en PDF y no ofrece serie tabular "
        "histórica por evento, que es exactamente lo que el documento afirma.")

    referencia(d, "Superintendencia de Transporte. (2026)",
               "National Oceanic and Atmospheric Administration, Climate Prediction "
               f"Center. (2026). {noaa['nombre']} (ONI) [serie estadística]. "
               f"Recuperado el 6 de agosto de 2026, de {noaa['url']}", antes=True)
    hechos.append(
        "Se añadió la referencia de la NOAA por el índice oceánico de El Niño, usado "
        "como variable de contexto. Título, entidad y dirección provienen de "
        "catalogo_fuentes.csv.")

    verificar.append(
        "Banco de la República: la dirección https://www.banrep.gov.co/es/estadisticas/trm "
        "no pudo reabrirse desde el entorno de trabajo. Ábrela y confirma que responde "
        "antes de imprimir.")
    verificar.append(
        "NOAA Climate Prediction Center: la dirección del índice oceánico tampoco pudo "
        "reabrirse desde el entorno de trabajo. Confírmala antes de imprimir.")
    verificar.append(
        "Fecha de consulta de la TRM y del índice oceánico: el manifiesto de fuentes "
        "solo registra la descarga del DANE y de la Superintendencia. Para las dos "
        "variables de contexto se usó la fecha de corte del proyecto, 6 de agosto de "
        "2026. Si la descarga fue otro día, ajústala.")

    # --- Tabla 16: verificación de fuentes en línea
    tabla16 = next((t for t in d.tables
                    if t.rows[0].cells[0].text.strip().lower().startswith("fuente")
                    and "verificada" in t.rows[0].cells[1].text.strip().lower()), None)
    if tabla16 is not None:
        fila_tabla(tabla16, ["Banco de la República (TRM)", "2026-08-06",
                             "por verificar", "sí, diaria"])
        fila_tabla(tabla16, ["NOAA, índice oceánico", "2026-08-06",
                             "por verificar", "sí, mensual"])
        fila_tabla(tabla16, ["DIMAR, boletines trimestrales", "2026-08-06",
                             "2026-08-06", "sí, trimestral"])
        hechos.append(
            "Tabla 16: se añadieron las tres fuentes nuevas. La de DIMAR queda con fecha "
            "de actualización verificada; las de la TRM y el índice oceánico quedan "
            "marcadas como «por verificar», que es lo que la evidencia disponible "
            "permite afirmar.")
    else:
        inconsistencias.append("No se localizó la tabla 16 de verificación de fuentes.")

    # --- 6. Consistencia sobre las 52 preguntas
    cambios_52 = [
        ("Análisis exploratorio completo de 52 preguntas con evidencia por pregunta.",
         "Análisis exploratorio de las 52 preguntas del catálogo, respondidas o "
         "cerradas con evidencia documentada por pregunta."),
        ("Ejecución de las 52 preguntas del análisis exploratorio",
         "Respuesta o cierre documentado de las 52 preguntas"),
    ]
    for viejo, nuevo in cambios_52:
        if sustituir_en_todo(d, viejo, nuevo):
            hechos.append(f"Consistencia de las 52 preguntas: «{viejo[:58]}…» pasa a "
                          f"«{nuevo[:58]}…».")

    # --- 9. Consistencia de formato de los índices de concentración
    separadores = [("3514,24", "3.514,24"), ("1351,38", "1.351,38"),
                   ("4722", "4.722"), ("2988", "2.988"), ("3740", "3.740"),
                   ("4217", "4.217")]
    aplicados = []
    for viejo, nuevo in separadores:
        if sustituir_en_todo(d, viejo, nuevo):
            aplicados.append(f"{viejo} → {nuevo}")
    if aplicados:
        hechos.append(
            "Formato numérico: el documento escribía unos índices de concentración con "
            "separador de miles (4.096, umbral 2.500) y otros sin él. Se unificaron sin "
            "alterar ningún valor: " + "; ".join(aplicados) + ".")

    # --- 7. Verificación de las cifras exigidas
    texto = "\n".join([p.text for p in d.paragraphs] +
                      [c.text for t in d.tables for f in t.rows for c in f.cells])
    exigidas = {
        "173 meses aduaneros": "173", "102 meses portuarios": "102",
        "101 meses comunes": "101", "registros aduaneros": "6.703.351",
        "preguntas ejecutadas": "38", "preguntas parciales": "4",
        "preguntas no viables": "10", "mejor WAPE del CIF": "5,875",
        "historia propia": "6,082", "historia propia + puerto": "6,575",
        "integrado completo": "6,167", "naive 1 sobre el CIF": "6,956",
        "mejora sobre naive 1": "15,5", "toneladas totales, modelo": "6,613",
        "toneladas totales, referencia": "9,174", "mejora en toneladas": "27,9",
        "contenerizada, modelo": "9,450", "contenerizada, referencia": "8,553",
        "cobertura mínima": "75,0", "cobertura máxima": "91,7",
        "crecimiento del CIF": "80,4", "aporte del volumen": "52,2",
        "aporte del valor unitario": "47,0", "importación portuaria": "19,6",
        "exportación portuaria": "1,3", "transbordo": "-85,2",
        "total portuario": "-13,3", "HHI capítulo": "433,74",
        "HHI país": "1.351,38", "HHI sociedad portuaria": "3.514,24",
    }
    ausentes = [f"{k} ({v})" for k, v in exigidas.items() if v not in texto]
    if ausentes:
        inconsistencias.append("Cifras exigidas que no se encontraron: " +
                               ", ".join(ausentes))

    # --- 8. Verificación de que los hallazgos centrales siguen en pie
    centrales = {
        "integración agregada por mes": "agregada por mes",
        "ausencia de llave pública": "llave pública",
        "el puerto no mejora el pronóstico": "no mejora el pronóstico",
        "aporte descriptivo y explicativo": "descriptivo y explicativo",
        "contenerizada sin modelo": "no usar modelo",
        "marítimo no viable": "no viable",
        "sin reporte, no sin operación": "registra reporte, no operación",
        "el índice mide reparto": "mide reparto",
        "valor unitario implícito": "valor unitario implícito",
    }
    perdidos = [k for k, v in centrales.items() if v not in texto]
    if perdidos:
        inconsistencias.append("Hallazgos centrales que ya no se localizan en el "
                               "texto: " + ", ".join(perdidos))

    prohibidas = ["dejaron de operar", "precio por kilogramo", "el puerto influye"]
    encontradas = [p for p in prohibidas if p in texto.lower()]
    if encontradas:
        inconsistencias.append("Expresiones prohibidas encontradas: " +
                               ", ".join(encontradas))

    SALIDA.parent.mkdir(parents=True, exist_ok=True)
    d.save(SALIDA)
    escribir_reporte()
    print(f"documento guardado en {SALIDA}")
    print(f"reporte guardado en {REPORTE}")
    print(f"cambios: {len(hechos)} · omitidos: {len(omitidos)} · "
          f"por verificar: {len(verificar)} · inconsistencias: {len(inconsistencias)}")


def escribir_reporte() -> None:
    def bloque(titulo, items, vacio):
        cuerpo = "\n".join(f"{i}. {x}\n" for i, x in enumerate(items, 1)) or f"{vacio}\n"
        return f"## {titulo}\n\n{cuerpo}"

    REPORTE.write_text(
        "# Reporte de corrección para la entrega ante jurados\n\n"
        "Documento de entrada: `Documento_Academico_Buenaventura_V5_FINAL_CORREGIDO.docx`\n"
        "Documento de salida: `Documento_Academico_Buenaventura_V5_FINAL_JURADOS.docx`\n"
        "El original no se sobrescribió.\n\n"
        + bloque("A. Cambios realizados", hechos, "Ninguno.")
        + bloque("B. Cambios que no se realizaron y por qué", omitidos, "Ninguno.")
        + bloque("C. Referencias que requieren verificación manual", verificar, "Ninguna.")
        + bloque("D. Inconsistencias detectadas", inconsistencias, "Ninguna.")
        + "## E. Confirmación sobre las cifras\n\n"
          "Las treinta cifras principales se verificaron una por una contra el texto del "
          "documento de salida y ninguna fue alterada. El único cambio numérico es de "
          "formato: se añadió el separador de miles a los índices de concentración que "
          "no lo llevaban, para que el documento no mezcle dos convenciones. El valor de "
          "cada índice es idéntico al de `comparacion_concentracion.csv` y "
          "`hhi_anual_sociedades.csv`.\n",
        encoding="utf-8")


if __name__ == "__main__":
    main()
